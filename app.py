import os
import logging
import time
import re
import json
import uuid
from datetime import datetime
from urllib.parse import quote_plus
from flask import Flask, render_template, request, jsonify, session, send_file, send_from_directory, url_for, make_response
from flask_sqlalchemy import SQLAlchemy
from sqlalchemy.orm import DeclarativeBase
from flask_login import LoginManager, login_user, current_user, login_required
from werkzeug.utils import secure_filename
import requests
import PyPDF2
from docx import Document
from ai_processor import process_text as legacy_process_text, chat_with_ai, split_text
from multi_provider_processor import multi_provider_processor
from PIL import Image
import pytesseract
import io
import base64
from io import BytesIO
from reportlab.pdfgen import canvas
from ai_detector import detect_ai_content # Added import
from humanizer import get_user_profile, add_writing_sample, clear_user_profile, get_user_style_text
import speech_recognition as sr
from pydub import AudioSegment
import tempfile
from style_rewrite_passthrough import process_style_rewrite
from coherence_service import (
    audit_job as run_coherence_audit,
    audit_claim as run_coherence_claim_audit,
    create_job as create_coherence_job,
    process_job_chunk as process_coherence_job_chunk,
    repair_next as repair_next_coherence_issue,
    skeleton_step as run_skeleton_step,
)

# Configure logging
logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

def convert_dollar_signs_to_text(text):
    """
    Convert dollar signs and monetary amounts to text format to prevent LaTeX formatting issues.
    Examples: $13 -> 13 dollars, $2,000 -> 2,000 dollars, $150.50 -> 150.50 dollars
    """
    if not text:
        return text
    
    # Pattern to match dollar amounts like $13, $2,000, $150.50, etc.
    # This matches: $ followed by digits, optional commas within digits, optional decimal with digits
    dollar_pattern = r'\$([0-9]{1,3}(?:,[0-9]{3})*(?:\.[0-9]{2})?)'
    
    # Replace dollar amounts with "X dollars" format
    converted_text = re.sub(dollar_pattern, r'\1 dollars', text)
    
    # Also handle standalone dollar signs that might be missed
    # Replace any remaining $ that aren't part of LaTeX math (avoid $$...$$)
    # Only replace single $ that are followed by space or at end of string
    converted_text = re.sub(r'\$(?!\$)(?=\s|$)', 'dollars', converted_text)
    
    return converted_text

UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'pdf', 'doc', 'docx', 'jpg', 'jpeg', 'png', 'mp3', 'wav'}
CHUNK_SIZE = 500  # Average words per chunk, approximately 1-2 pages

# Languages for translation
SUPPORTED_LANGUAGES = {
    'en': 'English',
    'es': 'Spanish (Español)',
    'fr': 'French (Français)',
    'de': 'German (Deutsch)',
    'it': 'Italian (Italiano)',
    'pt': 'Portuguese (Português)',
    'nl': 'Dutch (Nederlands)',
    'ru': 'Russian (Русский)',
    'zh': 'Chinese (中文)',
    'ja': 'Japanese (日本語)',
    'ko': 'Korean (한국어)',
    'ar': 'Arabic (العربية)',
    'hi': 'Hindi (हिन्दी)',
    'bn': 'Bengali (বাংলা)',
    'ur': 'Urdu (اردو)',
    'tr': 'Turkish (Türkçe)',
    'vi': 'Vietnamese (Tiếng Việt)',
    'th': 'Thai (ไทย)',
    'id': 'Indonesian (Bahasa Indonesia)',
    'pl': 'Polish (Polski)',
    'cs': 'Czech (Čeština)',
    'hu': 'Hungarian (Magyar)',
    'sv': 'Swedish (Svenska)',
    'fi': 'Finnish (Suomi)',
    'da': 'Danish (Dansk)',
    'no': 'Norwegian (Norsk)',
    'ro': 'Romanian (Română)',
    'uk': 'Ukrainian (Українська)',
    'he': 'Hebrew (עברית)',
    'fa': 'Persian (فارسی)',
    'el': 'Greek (Ελληνικά)'
}

# Languages supported by DeepL API
DEEPL_SUPPORTED_LANGUAGES = [
    'bg', 'cs', 'da', 'de', 'el', 'en', 'es', 'et', 'fi', 'fr', 'hu', 
    'id', 'it', 'ja', 'ko', 'lt', 'lv', 'nb', 'nl', 'pl', 'pt', 'ro', 
    'ru', 'sk', 'sl', 'sv', 'tr', 'uk', 'zh'
]

# Document length modes (in words)
SHORT_DOCUMENT_THRESHOLD = 2000
LONG_DOCUMENT_MAX = 500000

# Create uploads directory if it doesn't exist
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

class Base(DeclarativeBase):
    pass

db = SQLAlchemy(model_class=Base)
app = Flask(__name__)
app.secret_key = os.environ.get("SESSION_SECRET")
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 300 * 1024 * 1024

# Configure database connection using environment variables directly
app.config["SQLALCHEMY_DATABASE_URI"] = os.environ.get("NEON_DATABASE_URL") or os.environ.get("DATABASE_URL")
app.config["SQLALCHEMY_ENGINE_OPTIONS"] = {
    "pool_recycle": 300,
    "pool_pre_ping": True,
    "pool_size": 20,  # Increased pool size for better handling of concurrent connections
    "max_overflow": 10  # Allow up to 10 connections beyond pool_size
}

# Configure permanent sessions
app.config['PERMANENT_SESSION_LIFETIME'] = 86400  # 24 hours
app.config['SESSION_TYPE'] = 'filesystem'

# Initialize Flask-Login
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login'

db.init_app(app)

with app.app_context():
    try:
        import models
        db.create_all()
        logger.info("Successfully connected to database and created tables")
    except Exception as e:
        logger.error(f"Error initializing database: {str(e)}")
        logger.warning("App will continue without database connectivity")

@login_manager.user_loader
def load_user(user_id):
    return models.User.query.get(int(user_id))

@app.before_request
def before_request():
    if current_user.is_authenticated:
        session.permanent = True  # Make session permanent for logged in users

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def chunk_text(text):
    """
    Two-level chunking strategy:
    - For documents over 5,000 words: Split into macrochunks of ~5,000 words, each further split into subchunks of ~500 words
    - For smaller documents: Split directly into chunks of ~500 words
    """
    try:
        logger.debug(f"Starting text chunking, total length: {len(text)} characters")
        total_words = len(text.split())
        logger.debug(f"Total word count: {total_words}")
        
        # Determine chunking strategy based on document size
        MACROCHUNK_SIZE = 5000  # words (about 20-25 pages)
        
        # For documents over 5,000 words, use two-level chunking
        if total_words > MACROCHUNK_SIZE:
            logger.debug(f"Document exceeds {MACROCHUNK_SIZE} words, using two-level chunking strategy")
            chunks = two_level_chunking(text, MACROCHUNK_SIZE, CHUNK_SIZE)
        else:
            logger.debug(f"Document is under {MACROCHUNK_SIZE} words, using single-level chunking")
            chunks = single_level_chunking(text, CHUNK_SIZE)

        # Guarantee that every document over 2,000 words is divided into
        # multiple selectable sections, even if it contains one huge
        # paragraph or sentence with no natural splitting boundary.
        if total_words > SHORT_DOCUMENT_THRESHOLD and len(chunks) <= 1:
            words = text.split()
            chunks = [
                ' '.join(words[start:start + CHUNK_SIZE])
                for start in range(0, len(words), CHUNK_SIZE)
            ]
            logger.info(
                "Forced long document into %s selectable sections",
                len(chunks)
            )

        return chunks
        
    except Exception as e:
        logger.error(f"Error in chunking text: {str(e)}")
        # Fallback to simple chunking in case of error
        simple_chunks = [text[i:i+8000] for i in range(0, len(text), 8000)]
        logger.debug(f"Fallback: Created {len(simple_chunks)} simple chunks")
        return simple_chunks

def single_level_chunking(text, chunk_size):
    """
    Split text into manageable chunks of the specified size.
    
    Ensures that chunks start with complete sentences and maintains
    paragraph integrity where possible.
    """
    try:
        # Split text by paragraphs
        paragraphs = []
        # Handle different paragraph break styles
        for para_candidate in re.split(r'\n\s*\n', text):
            # Further split very long paragraphs
            if len(para_candidate.split()) > chunk_size * 1.5:  # Reduced threshold for better handling
                logger.debug(f"Breaking down very long paragraph: {len(para_candidate.split())} words")
                # Split by sentences for very long paragraphs
                sentences = re.split(r'(?<=[.!?])\s+', para_candidate)
                current_para = []
                current_length = 0
                
                for sentence in sentences:
                    sentence_words = len(sentence.split())
                    
                    # Special handling for extremely long sentences (edge case)
                    if sentence_words > chunk_size and not current_para:
                        # This one sentence exceeds our chunk size and would be the first sentence in the chunk
                        # We need to preserve it but mark it specifically
                        logger.warning(f"Found extremely long sentence: {sentence_words} words")
                        paragraphs.append(f"[LONG_SENTENCE]{sentence}")
                        continue
                        
                    if current_length + sentence_words > chunk_size and current_para:
                        paragraphs.append(' '.join(current_para))
                        current_para = [sentence]
                        current_length = sentence_words
                    else:
                        current_para.append(sentence)
                        current_length += sentence_words
                
                if current_para:
                    paragraphs.append(' '.join(current_para))
            else:
                if para_candidate.strip():
                    paragraphs.append(para_candidate.strip())
        
        logger.debug(f"Split text into {len(paragraphs)} paragraphs")
        
        # Now group paragraphs into chunks, being careful with sentence boundaries
        chunks = []
        current_chunk = []
        current_size = 0

        for paragraph in paragraphs:
            # Special handling for previously marked long sentences
            if paragraph.startswith("[LONG_SENTENCE]"):
                # If we have a current chunk in progress, finish it first
                if current_chunk:
                    chunks.append('\n\n'.join(current_chunk))
                    current_chunk = []
                    current_size = 0
                
                # Add this long sentence as its own chunk, removing the marker
                chunks.append(paragraph[15:])  # Remove the [LONG_SENTENCE] marker
                continue
                
            words = len(paragraph.split())
            
            # Create a new chunk if adding this paragraph would exceed the limit
            if current_size + words > chunk_size and current_chunk:
                chunks.append('\n\n'.join(current_chunk))
                current_chunk = [paragraph]
                current_size = words
            else:
                current_chunk.append(paragraph)
                current_size += words

        # Add the last chunk if there's anything left
        if current_chunk:
            chunks.append('\n\n'.join(current_chunk))
            
        logger.debug(f"Created {len(chunks)} chunks from text")

        # If no chunks were created, treat the entire text as one chunk
        result = chunks if chunks else [text]
        
        # Log some statistics
        if result:
            chunk_sizes = [len(chunk.split()) for chunk in result]
            logger.debug(f"Chunk sizes (words): min={min(chunk_sizes)}, max={max(chunk_sizes)}, avg={sum(chunk_sizes)/len(chunk_sizes):.1f}")
            
        return result
    except Exception as e:
        logger.error(f"Error in single-level chunking: {str(e)}")
        raise

def two_level_chunking(text, macrochunk_size, subchunk_size):
    """
    Split text into two levels:
    1. First into macrochunks of ~5,000 words
    2. Then each macrochunk into subchunks of ~500 words
    
    Ensures that macrochunks start with complete sentences.
    """
    try:
        # Split text into paragraphs
        paragraphs = []
        for para_candidate in re.split(r'\n\s*\n', text):
            if para_candidate.strip():
                paragraphs.append(para_candidate.strip())
        
        logger.debug(f"Split text into {len(paragraphs)} paragraphs for macrochunking")
        
        # Step 1: Group paragraphs into macrochunks
        macrochunks = []
        current_macrochunk = []
        current_size = 0

        for paragraph in paragraphs:
            words = len(paragraph.split())
            
            # For very long paragraphs, split them into sentences to ensure better chunking
            if words > macrochunk_size and not current_macrochunk:
                # This paragraph alone exceeds the macrochunk size and would be the first item
                # Split it into sentences to ensure complete sentence chunking
                sentences = re.split(r'(?<=[.!?])\s+', paragraph)
                current_sentence_group = []
                sentence_group_size = 0
                
                for sentence in sentences:
                    sentence_words = len(sentence.split())
                    if sentence_group_size + sentence_words > macrochunk_size and current_sentence_group:
                        # Complete the current macrochunk with sentences collected so far
                        macrochunks.append(' '.join(current_sentence_group))
                        current_sentence_group = [sentence]
                        sentence_group_size = sentence_words
                    else:
                        current_sentence_group.append(sentence)
                        sentence_group_size += sentence_words
                
                # Add any remaining sentences as their own macrochunk
                if current_sentence_group:
                    macrochunks.append(' '.join(current_sentence_group))
            
            # Normal case: Check if adding this paragraph would exceed the macrochunk size
            elif current_size + words > macrochunk_size and current_macrochunk:
                macrochunks.append('\n\n'.join(current_macrochunk))
                current_macrochunk = [paragraph]
                current_size = words
            else:
                current_macrochunk.append(paragraph)
                current_size += words

        # Add the last macrochunk if there's anything left
        if current_macrochunk:
            macrochunks.append('\n\n'.join(current_macrochunk))
            
        logger.debug(f"Created {len(macrochunks)} macrochunks from text")
        
        # Step 2: Split each macrochunk into subchunks
        all_subchunks = []
        for i, macrochunk in enumerate(macrochunks, 1):
            logger.debug(f"Processing macrochunk {i} of {len(macrochunks)}")
            subchunks = single_level_chunking(macrochunk, subchunk_size)
            all_subchunks.extend(subchunks)
        
        logger.debug(f"Created {len(all_subchunks)} total subchunks from {len(macrochunks)} macrochunks")
        
        # Log some statistics on the subchunks
        if all_subchunks:
            chunk_sizes = [len(chunk.split()) for chunk in all_subchunks]
            logger.debug(f"Subchunk sizes (words): min={min(chunk_sizes)}, max={max(chunk_sizes)}, avg={sum(chunk_sizes)/len(chunk_sizes):.1f}")
        
        return all_subchunks if all_subchunks else [text]
    except Exception as e:
        logger.error(f"Error in two-level chunking: {str(e)}")
        raise

def extract_text_from_pdf(file_path):
    """Extract text from PDF with better error handling and Unicode cleaning"""
    import re
    text = []
    try:
        logger.debug(f"Starting PDF extraction from {file_path}")
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page_num, page in enumerate(pdf_reader.pages, 1):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        # Clean Unicode surrogate characters and invalid characters
                        cleaned_text = re.sub(r'[\uD800-\uDFFF]', '', page_text)
                        cleaned_text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F-\x9F]', '', cleaned_text)
                        # Ensure UTF-8 compatibility
                        cleaned_text = cleaned_text.encode('utf-8', errors='ignore').decode('utf-8')
                        text.append(cleaned_text)
                    else:
                        logger.warning(f"No text extracted from page {page_num}")
                except Exception as page_error:
                    logger.error(f"Error extracting text from page {page_num}: {str(page_error)}")
                    continue

        if not text:
            raise ValueError("No text could be extracted from the PDF")

        return '\n\n'.join(text)
    except Exception as e:
        logger.error(f"Error in PDF extraction: {str(e)}")
        raise ValueError(f"Could not extract text from PDF: {str(e)}")

def extract_text_from_docx(file_path):
    """Extract text from DOCX with improved handling for complex documents"""
    try:
        logger.debug(f"Starting DOCX extraction from {file_path}")
        doc = Document(file_path)
        full_text = []
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text)
                if row_text:
                    full_text.append(' | '.join(row_text))
        
        # Handle headers and footers if available
        try:
            for section in doc.sections:
                # Try to access header and footer
                try:
                    if section.header:
                        for para in section.header.paragraphs:
                            if para.text.strip():
                                full_text.append(f"Header: {para.text}")
                except Exception as header_err:
                    logger.debug(f"Could not extract header: {str(header_err)}")
                
                try:
                    if section.footer:
                        for para in section.footer.paragraphs:
                            if para.text.strip():
                                full_text.append(f"Footer: {para.text}")
                except Exception as footer_err:
                    logger.debug(f"Could not extract footer: {str(footer_err)}")
        except Exception as section_err:
            logger.debug(f"Error processing sections: {str(section_err)}")
        
        # Join all text with double line breaks
        result = '\n\n'.join(full_text)
        logger.debug(f"Extracted {len(full_text)} text elements from DOCX")
        
        if not result.strip():
            raise ValueError("No text could be extracted from the document")
            
        return result
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {str(e)}")
        raise ValueError(f"Could not extract text from DOCX: {str(e)}")

def extract_text_from_image(image_path):
    """Extract text from an image using OCR"""
    try:
        image = Image.open(image_path)
        if image.mode != 'RGB':
            image = image.convert('RGB')
        text = pytesseract.image_to_string(image)
        return text.strip()
    except Exception as e:
        logger.error(f"Error extracting text from image: {str(e)}")
        raise

def extract_text_from_audio(audio_path):
    """Extract text from audio file using OpenAI's Whisper API"""
    try:
        logger.debug(f"Starting audio transcription from {audio_path}")
        
        # Check file exists
        if not os.path.exists(audio_path):
            raise ValueError(f"Audio file not found: {audio_path}")
        
        # Import the whisper transcription module
        from whisper_transcription import transcribe_audio_with_whisper
        
        # Use Whisper API for transcription
        text = transcribe_audio_with_whisper(audio_path)
        
        if not text or text.strip() == "":
            raise ValueError("Whisper API returned empty transcription")
            
        return text
    except ImportError:
        logger.error("Failed to import whisper_transcription module")
        raise ValueError("Audio transcription service is not properly configured")
    except Exception as e:
        logger.error(f"Error in Whisper audio transcription: {str(e)}")
        raise ValueError(f"Could not extract text from audio: {str(e)}")

def process_chunk(chunk_id, chunk_number, custom_instructions='', is_first_chunk=False, email=None, author_style='', content_source=None, ai_provider=None, preserve_length=True, style_source=None):
    """
    Process a single chunk with given instructions and optional user style
    using the multi-provider processor for improved reliability and performance
    
    Args:
        chunk_id: ID of the document containing the chunk
        chunk_number: Number of the chunk to process
        custom_instructions: Additional processing instructions
        is_first_chunk: Whether this is the first chunk (for style handling)
        email: User email for personal style
        author_style: Specific author style to imitate (e.g., "Orwell")
        content_source: Optional content source text to incorporate
        ai_provider: Optional AI provider to use (openai, anthropic, perplexity)
        preserve_length: Whether to maintain original text length (default: True)
    """
    start_time = time.time()
    
    # Get the chunk from the database
    chunk = models.DocumentChunk.query.filter_by(
        document_id=chunk_id,
        chunk_number=chunk_number
    ).first()

    if not chunk:
        raise ValueError('Chunk not found')
        
    # Update status to processing
    chunk.processing_status = 'processing'
    db.session.commit()

    try:
        # Get the document to check if it has a user profile
        document = models.TextEntry.query.get(chunk_id)
        
        # Get user style if available
        user_style_text = None
        if document and document.user_profile_id:
            user_style_text = get_user_style_text(profile_id=document.user_profile_id)
        elif email:
            user_style_text = get_user_style_text(email=email)
        if style_source:
            user_style_text = style_source
        
        # Get content source if available
        content_source_text = ""
        content_source_instructions = ""
        
        # First, check if a content source was provided directly via API
        if content_source:
            logger.debug(f"Using content source provided via API ({len(content_source)} characters)")
            content_source_text = content_source
        else:
            # If not, check for content sources in the database
            content_sources = models.ContentSource.query.filter_by(text_entry_id=chunk_id).all()
            if content_sources:
                logger.debug(f"Found {len(content_sources)} content sources for document {chunk_id}")
                for source in content_sources:
                    content_source_text += source.text_content + "\n\n"
                    if source.usage_instructions:
                        content_source_instructions += source.usage_instructions + "\n\n"
                
                logger.debug(f"Content source length from database: {len(content_source_text)}")
                if content_source_instructions:
                    logger.debug(f"Content source instructions: {content_source_instructions}")
        
        # If we have content source text from any source, log it
        if content_source_text:
            logger.debug(f"Total content source length: {len(content_source_text)} characters")
            
        # Check if this is the first chunk or a subsequent chunk
        include_style_in_output = is_first_chunk

        # Use the new multi-provider processor for better reliability
        logger.debug(f"Processing chunk {chunk_number} with multi-provider processor")
        
        # Generate effective instructions
        effective_instructions = custom_instructions
        
        # Add content source instructions if available
        if content_source_text:
            if not content_source_instructions:
                content_source_instructions = "Intelligently enrich the target document with relevant information, ideas, examples, and arguments from the content source, without overriding the target document's structure or identity."
            
            # Add content source instructions to custom instructions
            additional_instructions = (
                f"\n\nCONTENT SOURCE INSTRUCTIONS: {content_source_instructions}\n\n"
                f"CONTENT SOURCE TEXT:\n{content_source_text}\n\n"
                "Use the CONTENT SOURCE TEXT to enrich the Target Document (the input text) according to the CONTENT SOURCE INSTRUCTIONS."
            )
            
            if effective_instructions:
                effective_instructions += additional_instructions
            else:
                effective_instructions = additional_instructions
        
        # Add author style if specified
        if author_style:
            if effective_instructions:
                effective_instructions += f". Write in the style of {author_style}"
            else:
                effective_instructions = f"Write in the style of {author_style}"
            logger.debug(f"Added author style '{author_style}' to instructions")
        
        # Add mandatory length preservation instructions
        length_preservation_instructions = """
MANDATORY LENGTH PRESERVATION: Your rewritten output MUST match or exceed the length of the original text.
Target is 100-110% of the original word count.
Under NO circumstances should your output be shorter than the input.
"""
        if effective_instructions:
            effective_instructions = length_preservation_instructions + "\n" + effective_instructions
        else:
            effective_instructions = length_preservation_instructions
            
        try:
            # First, get the original word count
            original_word_count = len(chunk.original_chunk.split())
            logger.info(f"Original chunk word count: {original_word_count}")
            
            # Check if length multiplier is specified in custom instructions
            length_multiplier = None
            if '3x' in custom_instructions.lower() or 'triple' in custom_instructions.lower():
                length_multiplier = 3.0
            elif '2x' in custom_instructions.lower() or 'double' in custom_instructions.lower():
                length_multiplier = 2.0
            elif '1.5x' in custom_instructions.lower() or '150%' in custom_instructions.lower():
                length_multiplier = 1.5
                
            if length_multiplier:
                logger.info(f"Detected length multiplier: {length_multiplier}x in custom instructions")
            
            # Process the text - passing content_source directly to the processor rather than through instructions
            # Also pass the AI provider preference and maintain_length parameter
            provider_preference = None
            if ai_provider and ai_provider in ['openai', 'anthropic', 'perplexity']:
                provider_preference = ai_provider
                logger.info(f"Using user-selected AI provider: {provider_preference}")

            # Log content source usage for debugging
            if content_source_text:
                logger.info(f"Using content source with {len(content_source_text)} characters")
            
            processed_text = multi_provider_processor.process_text(
                text=chunk.original_chunk,
                action='rewrite',
                custom_instructions=effective_instructions,
                include_style_in_output=include_style_in_output,
                user_style_text=user_style_text,
                content_source=content_source_text if content_source_text else None,
                length_multiplier=length_multiplier,
                provider_preference=provider_preference,
                maintain_length=preserve_length
            )
            
            # Verify length requirements
            processed_word_count = len(processed_text.split())
            length_ratio = processed_word_count / original_word_count
            logger.info(f"Processed chunk word count: {processed_word_count}, ratio: {length_ratio:.2f}")
            
            # If output is shorter than input, force expansion
            max_expansion_attempts = 3
            current_attempt = 0
            
            while length_ratio < 1.0 and current_attempt < max_expansion_attempts:
                current_attempt += 1
                logger.warning(f"Output too short ({length_ratio:.2f}). Emergency expansion attempt {current_attempt}/{max_expansion_attempts}")
                
                emergency_instructions = f"""
EMERGENCY EXPANSION REQUIRED: Your rewrite is too short!
Original text: {original_word_count} words
Your rewrite: {processed_word_count} words

You MUST expand the text to AT LEAST {original_word_count} words, preferably {int(original_word_count * 1.1)} words.
Expand by:
1. Adding detailed examples and evidence for each point
2. Elaborating on existing concepts with clarifications
3. Providing deeper explanations and implications
4. DO NOT add unrelated material or irrelevant content
5. DO NOT summarize - this is a rewrite with expansion

Your expanded text MUST preserve the intellectual depth and argumentative structure of the original.
"""
                
                processed_text = multi_provider_processor.process_text(
                    text=processed_text,
                    action='rewrite',
                    custom_instructions=final_emergency_instructions,
                    include_style_in_output=False,
                    user_style_text=user_style_text,
                    content_source=content_source_text if content_source_text else None,
                    length_multiplier=length_multiplier
                )
                
                # Convert dollar signs to "dollars" after final expansion
                processed_text = convert_dollar_signs_to_text(processed_text)
                
                # Re-check length requirements
                processed_word_count = len(processed_text.split())
                length_ratio = processed_word_count / original_word_count
                logger.info(f"After expansion attempt {current_attempt}: {processed_word_count} words, ratio: {length_ratio:.2f}")
                
                if length_ratio >= 1.0:
                    logger.info(f"Length requirement met after {current_attempt} expansion attempts")
                    break
            
            # If still too short after all attempts, make one final aggressive attempt
            if length_ratio < 1.0:
                logger.warning(f"Still below minimum length ({length_ratio:.2f}) after {max_expansion_attempts} attempts. Final attempt.")
                
                final_emergency_instructions = f"""
CRITICAL LENGTH FAILURE - FINAL ATTEMPT:
Your output is MUCH TOO SHORT. The original text had {original_word_count} words, but your rewrite has only {processed_word_count} words.

You MUST expand this text to AT LEAST {original_word_count} words, preferably {int(original_word_count * 1.1)} words.
This is NON-NEGOTIABLE - your primary task is to meet the length requirement.

For EACH paragraph:
1. DOUBLE the size with detailed examples, evidence, and elaboration
2. Add substantial explanations for each concept
3. Explore the implications and applications of each idea
4. Maintain all intellectual arguments and flow from the original

DO NOT summarize. DO NOT condense. EXPAND with meaningful content.
"""
                
                processed_text = multi_provider_processor.process_text(
                    text=processed_text,
                    action='rewrite',
                    custom_instructions=final_emergency_instructions,
                    include_style_in_output=False,
                    user_style_text=user_style_text,
                    content_source=content_source_text if content_source_text else None,
                    length_multiplier=length_multiplier
                )
                
                # Convert dollar signs to "dollars" after final expansion
                processed_text = convert_dollar_signs_to_text(processed_text)
                
                processed_word_count = len(processed_text.split())
                length_ratio = processed_word_count / original_word_count
                logger.info(f"After final expansion attempt: {processed_word_count} words, ratio: {length_ratio:.2f}")
            
        except Exception as e:
            logger.error(f"Error with multi-provider processor: {str(e)}. Falling back to legacy processor.")
            # Fallback to legacy processor if the new processor fails
            processed_text = legacy_process_text(
                text=chunk.original_chunk,
                action='rewrite',
                custom_instructions=effective_instructions if 'effective_instructions' in locals() else custom_instructions,
                include_style_in_output=include_style_in_output,
                user_style_text=user_style_text
            )
            
            # Convert dollar signs to "dollars" after legacy processing
            processed_text = convert_dollar_signs_to_text(processed_text)
            
            # Even with legacy processor, verify length requirements
            original_word_count = len(chunk.original_chunk.split())
            processed_word_count = len(processed_text.split())
            length_ratio = processed_word_count / original_word_count
            
            if length_ratio < 1.0:
                logger.warning(f"Legacy processor output too short ({length_ratio:.2f}). Adding emergency expansion instructions.")
                emergency_instructions = f"EXPAND this text to at least {original_word_count} words while preserving all meaning and intellectual depth."
                
                try:
                    processed_text = legacy_process_text(
                        text=processed_text,
                        action='expand',
                        custom_instructions=emergency_instructions,
                        include_style_in_output=False,
                        user_style_text=user_style_text
                    )
                    
                    # Convert dollar signs to "dollars" after legacy expansion
                    processed_text = convert_dollar_signs_to_text(processed_text)
                except Exception as expand_error:
                    logger.error(f"Error expanding with legacy processor: {str(expand_error)}")
                    # Continue with the best we have at this point

        # Check for timeout
        if time.time() - start_time > 45:  # 45-second timeout
            raise TimeoutError(f"Processing chunk {chunk_number} timed out")
            
        # Convert dollar signs to "dollars" to prevent LaTeX formatting issues
        processed_text = convert_dollar_signs_to_text(processed_text)
        
        # Translate if target language is specified
        if document and document.target_language and document.target_language != 'en':
            logger.debug(f"Translating to {document.target_language}")
            try:
                # Track the original length for validation
                original_word_count = len(processed_text.split())
                logger.debug(f"Original text word count before translation: {original_word_count}")
                
                # Use our enhanced chunking translation
                translated_text = translate_text(processed_text, document.target_language)
                processed_text = translated_text
                
                # Validate the translation was complete
                translated_word_count = len(processed_text.split())
                logger.debug(f"Translation completed, length: {len(processed_text)}, word count: {translated_word_count}")
                
                # Verify we didn't lose a significant amount of content
                if translated_word_count < original_word_count * 0.6 and translated_word_count < 100:
                    logger.error(f"Translation validation failed - output too short: {translated_word_count}/{original_word_count} words")
                    # Send alert to logging but continue with what we have
            except Exception as translation_error:
                logger.error(f"Translation error: {str(translation_error)}")
                # Continue with untranslated text if translation fails

        chunk.processed_chunk = processed_text
        chunk.is_processed = True
        chunk.processing_status = 'complete'
        db.session.commit()
        
        logger.debug(f"Successfully processed chunk {chunk_number} in {time.time() - start_time:.2f} seconds")
        
        return chunk
        
    except Exception as e:
        # Get detailed error information
        import traceback
        error_traceback = traceback.format_exc()
        
        # Mark the chunk as having an error
        chunk.processing_status = 'error'
        chunk.processed_chunk = f"[Error processing chunk: {str(e)}]"
        db.session.commit()
        
        # Log both the error message and full traceback
        logger.error(f"Error processing chunk {chunk_number}: {str(e)}")
        logger.error(f"Full traceback for chunk {chunk_number}: {error_traceback}")
        
        # Re-raise the exception to be handled by the calling function
        raise

@app.route('/')
def index():
    return render_template('index.html')
    
@app.route('/translate')
def translation_page():
    """Render the dedicated translation page with source and target language boxes"""
    return render_template('translation_page.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400

        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400

        if not allowed_file(file.filename):
            return jsonify({'error': 'File type not supported'}), 400

        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)

        try:
            # Extract text based on file type
            if filename.lower().endswith('.pdf'):
                text = extract_text_from_pdf(file_path)
            elif filename.lower().endswith(('.doc', '.docx')):
                text = extract_text_from_docx(file_path)
            elif filename.lower().endswith(('.jpg', '.jpeg', '.png')):
                text = extract_text_from_image(file_path)
            elif filename.lower().endswith(('.mp3', '.wav')):
                text = extract_text_from_audio(file_path)
            else:
                return jsonify({'error': 'Unsupported file format'}), 400

            # Split text into chunks
            chunks = chunk_text(text)
            if not chunks:
                return jsonify({'error': 'No text could be extracted'}), 400

            logger.debug(f"Created {len(chunks)} chunks from the text")

            # Create text entry and chunks
            text_entry = models.TextEntry(
                original_text=text,
                processed_text="",  # Will be built from chunks
                action="rewrite",
                complexity="default",
                total_chunks=len(chunks)
            )
            db.session.add(text_entry)
            db.session.flush()

            # Create chunks
            for i, chunk in enumerate(chunks, 1):
                doc_chunk = models.DocumentChunk(
                    document_id=text_entry.id,
                    chunk_number=i,
                    original_chunk=chunk,
                    processed_chunk="",
                    processing_status = "pending"
                )
                db.session.add(doc_chunk)

            db.session.commit()

            # Store document ID in session
            session['current_document_id'] = text_entry.id
            session['current_chunk'] = 1
            
            # Calculate document stats
            word_count = len(text.split())
            chars_count = len(text)
            ave_chunk_size = word_count // len(chunks) if len(chunks) > 0 else 0

            # Send the ENTIRE text, not just the first chunk
            return jsonify({
                'text': text,  # Send the complete original text
                'document_id': text_entry.id,
                'total_chunks': len(chunks),
                'current_chunk': 1,
                'file_name': filename,
                'total_words': word_count,
                'total_chars': chars_count,
                'average_chunk_size': ave_chunk_size,
                'full_document_loaded': True  # Flag to indicate the full document was processed
            })

        except Exception as e:
            logger.error(f"Error processing file: {str(e)}")
            return jsonify({'error': f'Error processing file: {str(e)}'}), 500
        finally:
            if os.path.exists(file_path):
                os.remove(file_path)

    except Exception as e:
        logger.error(f"Error in file upload: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/get_chunk', methods=['POST', 'GET'])
def get_chunk():
    try:
        # Support both GET and POST methods
        if request.method == 'GET':
            document_id = request.args.get('document_id')
            chunk_number = request.args.get('chunk_number', 1, type=int)
            get_all = request.args.get('all', 'false').lower() == 'true'
        else:
            data = request.get_json()
            document_id = data.get('document_id')
            chunk_number = data.get('chunk_number', 1)
            get_all = data.get('all', False)

        if not document_id:
            return jsonify({'error': 'No document selected'}), 400
            
        # Get the document
        document = models.TextEntry.query.get(document_id)
        if not document:
            return jsonify({'error': 'Document not found'}), 404
            
        # If 'all' parameter is true, return all chunks
        if get_all:
            logger.debug(f"Retrieving all chunks for document {document_id}")
            chunks = models.DocumentChunk.query.filter_by(
                document_id=document_id
            ).order_by(models.DocumentChunk.chunk_number).all()
            
            if not chunks:
                return jsonify({'error': 'No chunks found for document'}), 404
                
            return jsonify({
                'chunks': [chunk.to_dict() for chunk in chunks],
                'total_chunks': document.total_chunks
            })
        else:
            # Return a single chunk
            chunk = models.DocumentChunk.query.filter_by(
                document_id=document_id,
                chunk_number=chunk_number
            ).first()

            if not chunk:
                return jsonify({'error': 'Chunk not found'}), 404

            return jsonify({
                'chunk': chunk.to_dict(),
                'total_chunks': document.total_chunks
            })

    except Exception as e:
        logger.error(f"Error getting chunk: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/process_chunk', methods=['POST'])
def process_chunk_route():
    try:
        data = request.get_json()
        document_id = data.get('document_id')
        chunk_number = data.get('chunk_number')
        custom_instructions = data.get('custom_instructions', '')
        is_first_chunk = data.get('is_first_chunk', False)
        email = data.get('email', '')
        author_style = data.get('author_style', '')
        target_language = data.get('target_language', '')
        content_source = data.get('content_source', '')  # Get content source from the request
        ai_provider = data.get('ai_provider', '')  # Get user's AI provider preference
        preserve_length = data.get('preserve_length', True)  # Get length preservation setting
        style_source = data.get('style_source', '')

        if not document_id or not chunk_number:
            return jsonify({'error': 'Missing document_id or chunk_number'}), 400

        # If no email provided, try to get from session
        if not email:
            email = session.get('last_email')
            
        # Get the document to save target language if provided
        if target_language:
            document = models.TextEntry.query.get(document_id)
            if document:
                document.target_language = target_language
                db.session.commit()
                logger.debug(f"Set target language for document {document_id} to {target_language}")
        
        # If content_source is provided in the request, add it to the instructions
        content_source_instructions = ''
        if content_source:
            # Log the length to avoid logging the entire content source which could be large
            logger.debug(f"Received content source in request: {len(content_source)} characters")
            
            # Add a specific marker in the instructions for content source if not already included
            if 'content source' not in custom_instructions.lower() and 'use the provided content' not in custom_instructions.lower():
                content_source_instructions = " Use the provided content source to enrich the text."
            
            # Add the content source instruction to custom_instructions
            custom_instructions = custom_instructions.strip() + content_source_instructions

        try:
            # Pass all parameters to the process_chunk function including content_source, ai_provider, and preserve_length
            chunk = process_chunk(
                document_id, 
                chunk_number, 
                custom_instructions, 
                is_first_chunk, 
                email,
                author_style,  # Add author style parameter
                content_source,  # Pass content source to process_chunk
                ai_provider,  # Pass the selected AI provider
                preserve_length,  # Pass length preservation setting
                style_source
            )
            return jsonify({'chunk': chunk.to_dict()})
        except ValueError as e:
            return jsonify({'error': str(e)}), 404

    except Exception as e:
        import traceback
        error_traceback = traceback.format_exc()
        logger.error(f"Error in process_chunk_route: {str(e)}")
        logger.error(f"Traceback in process_chunk_route: {error_traceback}")
        return jsonify({'error': str(e), 'traceback': error_traceback}), 500

@app.route('/process_all_chunks', methods=['POST'])
def process_all_chunks():
    try:
        data = request.get_json()
        document_id = data.get('document_id')
        custom_instructions = data.get('custom_instructions', '')
        # Add support for author style
        author_style = data.get('author_style', '')
        email = data.get('email', '')
        target_language = data.get('target_language', '')

        if not document_id:
            return jsonify({'error': 'Missing document_id'}), 400
            
        # Get the document to save target language if provided
        if target_language:
            document = models.TextEntry.query.get(document_id)
            if document:
                document.target_language = target_language
                db.session.commit()
                logger.debug(f"Set target language for document {document_id} to {target_language}")

        # Get all chunks for the document
        chunks = models.DocumentChunk.query.filter_by(
            document_id=document_id
        ).order_by(models.DocumentChunk.chunk_number).all()

        if not chunks:
            return jsonify({'error': 'No chunks found'}), 404

        total_chunks = len(chunks)
        processed_chunks = 0
        errors = []

        # Find the next unprocessed chunk
        current_chunk = None
        for chunk in chunks:
            if not chunk.is_processed:
                current_chunk = chunk
                break

        if not current_chunk:
            # All chunks are processed
            return jsonify({
                'status': 'complete',
                'processed_chunks': total_chunks,
                'total_chunks': total_chunks,
                'percentage': 100
            })

        try:
            # Process the current chunk with timeout
            start_time = time.time()

            # Mark chunk as processing
            current_chunk.processing_status = "processing"
            db.session.commit()

            # Determine if this is the first chunk - only chunk #1 should include style instructions
            is_first_chunk = current_chunk.chunk_number == 1
            
            logger.debug(f"Processing chunk {current_chunk.chunk_number}, is_first_chunk={is_first_chunk}")

            # Get the document to check if it has a user profile
            document = models.TextEntry.query.get(document_id)
            
            # Get user style if available
            user_style_text = None
            if document and document.user_profile_id:
                user_style_text = get_user_style_text(profile_id=document.user_profile_id)
            
            # Use email from session if available
            email = session.get('last_email')
            if not user_style_text and email:
                user_style_text = get_user_style_text(email=email)

            # Use the new multi-provider processor for better reliability
            try:
                # If author style is specified, include it in custom instructions with maximum emphasis
                effective_instructions = custom_instructions
                if author_style:
                    # CRITICAL: This must be the first instruction and given highest priority
                    author_style_instruction = f"!!! ABSOLUTELY MANDATORY !!! STRICTLY WRITE IN THE EXACT STYLE OF {author_style.upper()} - THIS IS THE HIGHEST PRIORITY INSTRUCTION AND MUST BE FOLLOWED EXACTLY !!!"
                    
                    if effective_instructions:
                        effective_instructions = author_style_instruction + "\n\n" + effective_instructions
                    else:
                        effective_instructions = author_style_instruction
                # Pass user instructions without automatic modifications
                # No automatic length or conciseness instructions that override user requirements
                    
                # First, get the original word count
                original_word_count = len(current_chunk.original_chunk.split())
                logger.info(f"Original chunk word count: {original_word_count}")
                
                # Determine if user selected a specific AI provider
                provider_preference = None
                if ai_provider and ai_provider in ['openai', 'anthropic', 'perplexity']:
                    provider_preference = ai_provider
                    logger.info(f"Using user-selected AI provider: {provider_preference}")
                
                processed_text = multi_provider_processor.process_text(
                    text=current_chunk.original_chunk,
                    action='rewrite',
                    custom_instructions=effective_instructions,
                    include_style_in_output=is_first_chunk,
                    user_style_text=user_style_text,
                    provider_preference=provider_preference,
                    maintain_length=preserve_length
                )
                
                # Convert dollar signs to "dollars" to prevent LaTeX formatting issues
                processed_text = convert_dollar_signs_to_text(processed_text)
                
                # Verify length requirements
                processed_word_count = len(processed_text.split())
                length_ratio = processed_word_count / original_word_count
                logger.info(f"Processed chunk word count: {processed_word_count}, ratio: {length_ratio:.2f}")
                
                # If output is shorter than input, force expansion
                if length_ratio < 1.0:
                    logger.warning(f"Output too short ({length_ratio:.2f}). Performing emergency expansion.")
                    
                    emergency_instructions = f"""
EMERGENCY EXPANSION REQUIRED: Your rewrite is too short!
Original text: {original_word_count} words
Your rewrite: {processed_word_count} words

You MUST expand the text to AT LEAST {original_word_count} words, preferably {int(original_word_count * 1.1)} words.
This is NON-NEGOTIABLE - length preservation is the primary requirement.

Expand by:
1. Adding detailed examples and evidence for each point
2. Elaborating on existing concepts with clarifications
3. Providing deeper explanations and implications

Your expanded text MUST preserve the intellectual depth and argumentative structure of the original.
"""
                    
                    # Try to expand the text
                    processed_text = multi_provider_processor.process_text(
                        text=processed_text,
                        action='rewrite',
                        custom_instructions=emergency_instructions,
                        include_style_in_output=False,
                        user_style_text=user_style_text
                    )
                    
                    # Convert dollar signs to "dollars" after expansion
                    processed_text = convert_dollar_signs_to_text(processed_text)
                    
                    # Re-check length
                    processed_word_count = len(processed_text.split())
                    length_ratio = processed_word_count / original_word_count
                    logger.info(f"After expansion: {processed_word_count} words, ratio: {length_ratio:.2f}")
                
            except Exception as e:
                logger.error(f"Error with multi-provider processor: {str(e)}. Falling back to legacy processor.")
                # Fallback to legacy processor if the new processor fails
                processed_text = legacy_process_text(
                    text=current_chunk.original_chunk,
                    action='rewrite',
                    custom_instructions=effective_instructions if 'effective_instructions' in locals() else custom_instructions,
                    include_style_in_output=is_first_chunk,
                    user_style_text=user_style_text
                )
                
                # Even with legacy processor, verify length requirements
                if 'original_word_count' not in locals():
                    original_word_count = len(current_chunk.original_chunk.split())
                
                processed_word_count = len(processed_text.split())
                length_ratio = processed_word_count / original_word_count
                
                if length_ratio < 1.0:
                    logger.warning(f"Legacy processor output too short ({length_ratio:.2f}). Adding emergency expansion.")
                    try:
                        emergency_instructions = f"EXPAND this text to at least {original_word_count} words while preserving meaning."
                        
                        processed_text = legacy_process_text(
                            text=processed_text,
                            action='expand',
                            custom_instructions=emergency_instructions,
                            include_style_in_output=False
                        )
                    except Exception as expand_error:
                        logger.error(f"Error expanding with legacy processor: {str(expand_error)}")

            # Check for timeout
            if time.time() - start_time > 30:  # Increased timeout to 30 seconds
                raise TimeoutError(f"Processing chunk {current_chunk.chunk_number} timed out")

            if "Error processing text:" in processed_text:
                raise Exception(processed_text)

            # Update chunk status
            current_chunk.processed_chunk = processed_text
            current_chunk.is_processed = True
            current_chunk.processing_status = "complete"
            db.session.commit()

            # Count actually processed chunks
            processed_chunks = models.DocumentChunk.query.filter_by(
                document_id=document_id,
                is_processed=True
            ).count()

            # Return progress
            percentage = int((processed_chunks / total_chunks) * 100)
            return jsonify({
                'status': 'in_progress',
                'processed_chunks': processed_chunks,
                'total_chunks': total_chunks,
                'percentage': percentage,
                'current_chunk': current_chunk.chunk_number,
                'processed_text': processed_text,
                'errors': errors if errors else None
            })

        except Exception as chunk_error:
            error_msg = str(chunk_error)
            logger.error(f"Error processing chunk {current_chunk.chunk_number}: {error_msg}")

            # Mark chunk for retry
            current_chunk.processing_status = "error"
            current_chunk.is_processed = False  # Reset for retry
            db.session.commit()

            errors.append(f"Chunk {current_chunk.chunk_number}: {error_msg}")

            # Return error status
            return jsonify({
                'error': error_msg,
                'processed_chunks': processed_chunks,
                'total_chunks': total_chunks,
                'percentage': int((processed_chunks / total_chunks) * 100),
                'current_chunk': current_chunk.chunk_number,
                'status': 'error',
                'retry': True
            }), 500

    except Exception as e:
        logger.error(f"Error in process_all_chunks: {str(e)}")
        return jsonify({
            'error': str(e),
            'processed_chunks': processed_chunks if 'processed_chunks' in locals() else 0,
            'total_chunks': total_chunks if 'total_chunks' in locals() else 0,
            'percentage': int((processed_chunks / total_chunks) * 100) if 'processed_chunks' in locals() and 'total_chunks' in locals() and total_chunks > 0 else 0
        }), 500


@app.route('/api/coherence/start', methods=['POST'])
def start_coherence_job():
    """Create a resumable document-scale coherence job."""
    try:
        if os.environ.get('COHERENCE_ENABLED', 'true').lower() not in ('1', 'true', 'yes', 'on'):
            return jsonify({
                'error': 'Document-scale coherence is currently disabled',
                'code': 'coherence_disabled',
            }), 409
        data = request.get_json() or {}
        document_id = data.get('document_id')
        document = models.TextEntry.query.get(document_id)
        if not document:
            return jsonify({'error': 'Document not found'}), 404
        if int(session.get('current_document_id', -1)) != int(document_id):
            return jsonify({'error': 'This document is not active in your session'}), 403
        if len(document.original_text.split()) <= SHORT_DOCUMENT_THRESHOLD:
            return jsonify({'error': 'Document-scale coherence is for documents over 2,000 words'}), 400

        job = create_coherence_job(
            db.session,
            document,
            custom_instructions=data.get('custom_instructions', ''),
            author_style=data.get('author_style', ''),
            style_source=(
                data.get('style_source', '')
                or (get_user_style_text(data.get('email')) if data.get('email') else '')
            ),
            content_source=data.get('content_source', ''),
            provider_preference=data.get('ai_provider', ''),
            selected_chunks=data.get('selected_chunks'),
        )
        session['coherence_job_id'] = job.id
        known_jobs = session.get('coherence_job_ids', [])
        session['coherence_job_ids'] = list(dict.fromkeys(
            [*known_jobs, job.id]
        ))[-20:]
        return jsonify({
            'success': True,
            'job_id': job.id,
            'status': job.status,
            'total_chunks': job.num_chunks,
            'length_mode': job.length_mode,
            'target_words': {
                'minimum': job.target_min_words,
                'maximum': job.target_max_words,
                'target': job.target_mid_words,
            },
        })
    except Exception as error:
        db.session.rollback()
        logger.exception("Unable to initialize coherence job")
        return jsonify({'error': str(error)}), 500


@app.route('/api/coherence/<int:job_id>/skeleton-step', methods=['POST'])
def coherence_skeleton_step(job_id):
    """Analyze one source group and progressively build the frozen skeleton."""
    try:
        if job_id not in session.get('coherence_job_ids', []):
            return jsonify({'error': 'Coherence job not found in this session'}), 403
        result = run_skeleton_step(db.session, multi_provider_processor, job_id)
        return jsonify({'success': True, **result})
    except Exception as error:
        db.session.rollback()
        logger.exception("Coherence skeleton extraction failed")
        return jsonify({'error': str(error)}), 500


@app.route('/api/coherence/<int:job_id>/process/<int:chunk_index>', methods=['POST'])
def coherence_process_chunk(job_id, chunk_index):
    """Process one section against the frozen skeleton and live memory."""
    try:
        if job_id not in session.get('coherence_job_ids', []):
            return jsonify({'error': 'Coherence job not found in this session'}), 403
        state = process_coherence_job_chunk(
            db.session,
            multi_provider_processor,
            job_id,
            chunk_index,
            convert_dollar_signs_to_text,
        )
        return jsonify({
            'success': True,
            'chunk_index': state.chunk_index,
            'processed_text': state.chunk_output_text,
            'actual_words': state.actual_words,
            'delta': state.chunk_delta,
        })
    except Exception as error:
        db.session.rollback()
        logger.exception("Coherent section processing failed")
        return jsonify({'error': str(error)}), 500


@app.route('/api/coherence/<int:job_id>/audit', methods=['POST'])
def coherence_audit(job_id):
    """Run the global consistency stitch and create a micro-repair plan."""
    try:
        if job_id not in session.get('coherence_job_ids', []):
            return jsonify({'error': 'Coherence job not found in this session'}), 403
        report = run_coherence_audit(
            db.session, multi_provider_processor, job_id
        )
        job = models.CoherenceJob.query.get(job_id)
        return jsonify({
            'success': True,
            'report': report,
            'repairs_required': len(report.get('conflicts', [])),
            'complete': job.status == 'complete',
            'output': job.final_output if job.status == 'complete' else None,
            'word_count': job.final_word_count if job.status == 'complete' else None,
        })
    except Exception as error:
        db.session.rollback()
        logger.exception("Coherence audit failed")
        return jsonify({'error': str(error)}), 500


@app.route('/api/coherence/<int:job_id>/repair-next', methods=['POST'])
def coherence_repair_next(job_id):
    """Execute one targeted repair from the global consistency plan."""
    try:
        if job_id not in session.get('coherence_job_ids', []):
            return jsonify({'error': 'Coherence job not found in this session'}), 403
        result = repair_next_coherence_issue(
            db.session,
            multi_provider_processor,
            job_id,
            convert_dollar_signs_to_text,
        )
        return jsonify({'success': True, **result})
    except Exception as error:
        db.session.rollback()
        logger.exception("Coherence repair failed")
        return jsonify({'error': str(error)}), 500


@app.route('/api/coherence/<int:job_id>/status', methods=['GET'])
def coherence_job_status(job_id):
    """Expose persisted phase/progress state for inspection and resumption."""
    if job_id not in session.get('coherence_job_ids', []):
        return jsonify({'error': 'Coherence job not found in this session'}), 403
    job = models.CoherenceJob.query.get(job_id)
    if not job:
        return jsonify({'error': 'Coherence job not found'}), 404
    chunks = models.CoherenceChunk.query.filter_by(job_id=job_id).order_by(
        models.CoherenceChunk.chunk_index
    ).all()
    return jsonify({
        'success': True,
        'job_id': job.id,
        'status': job.status,
        'skeleton_ready': bool(job.global_skeleton) and job.status != 'skeleton_extraction',
        'skeleton_steps_completed': job.skeleton_cursor,
        'current_chunk': job.current_chunk,
        'total_chunks': job.num_chunks,
        'completed_chunks': sum(1 for chunk in chunks if chunk.status == 'complete'),
        'completed_chunk_numbers': [
            chunk.chunk_index for chunk in chunks if chunk.status == 'complete'
        ],
        'chunks': [
            {
                'chunk_index': chunk.chunk_index,
                'status': chunk.status,
                'processed_text': chunk.chunk_output_text
                if chunk.status == 'complete' else None,
            }
            for chunk in chunks
        ],
        'failed_chunks': [
            {'chunk_index': chunk.chunk_index, 'error': chunk.error_message}
            for chunk in chunks if chunk.status == 'failed'
        ],
        'validation_report': job.validation_report,
        'warnings': job.warnings or [],
        'output': job.final_output if job.status == 'complete' else None,
        'word_count': job.final_word_count,
        'error': job.error_message,
    })


@app.route('/api/coherence/<int:job_id>/claim-audit', methods=['POST'])
def coherence_claim_audit(job_id):
    """Verify claims at any point against the job's current Tractatus memory."""
    try:
        if job_id not in session.get('coherence_job_ids', []):
            return jsonify({'error': 'Coherence job not found in this session'}), 403
        data = request.get_json() or {}
        report = run_coherence_claim_audit(
            db.session,
            job_id,
            text=data.get('text'),
            chunk_index=data.get('chunk_index'),
        )
        return jsonify({'success': True, 'report': report})
    except Exception as error:
        logger.exception("Claim audit failed")
        return jsonify({'error': str(error)}), 400

@app.route('/process', methods=['POST'])
def process():
    try:
        data = request.get_json()
        if not data:
            return jsonify({'error': 'No data provided'}), 400

        text = data.get('text', '')
        custom_instructions = data.get('custom_instructions', '')
        email = data.get('email', '')
        author_style = data.get('author_style', '')
        ai_provider = data.get('ai_provider', '')
        preserve_length = data.get('preserve_length', True)
        target_language = data.get('target_language', '')
        content_source = data.get('content_source', '')
        style_source = data.get('style_source', '')
        processing_mode = data.get('processing_mode', 'rewrite')

        if not text:
            return jsonify({'error': 'No text provided'}), 400

        # Initialize user_style_text for all modes
        user_style_text = None
        if email:
            user_style_text = get_user_style_text(email=email)
        else:
            # Try to get from session
            email = session.get('last_email')
            if email:
                user_style_text = get_user_style_text(email=email)
        if style_source:
            user_style_text = style_source

        # Handle different processing modes
        if processing_mode == 'homework':
            # Homework mode: treat input text as instructions to follow
            instructions_to_follow = text
            
            # Build homework completion prompt
            homework_prompt = f"""You are an AI assistant helping to complete a task, assignment, or follow instructions. The following text contains the instructions or questions you need to address:

INSTRUCTIONS/TASK:
{instructions_to_follow}

Please complete this task thoroughly and professionally. Follow all instructions exactly as given. If it's an exam, answer all questions completely. If it's homework, solve all problems with explanations. If it's a set of instructions, follow them precisely.

Additional guidance: {custom_instructions if custom_instructions else 'None provided.'}
"""
            
            # For homework mode, we use the prompt as the text to process
            text = homework_prompt
            
            # Clear custom instructions since they're now part of the main prompt
            custom_instructions = ""
            
            # Homework mode doesn't need length preservation or author style
            preserve_length = False
            author_style = ""
            
            logger.info(f"Processing in homework mode with {len(instructions_to_follow)} characters of instructions")
        else:
            # Rewrite mode: traditional text transformation
            # Add author style to custom instructions if provided, with MAXIMUM emphasis
            if author_style:
                # CRITICAL: This must be the first instruction and given highest priority
                author_style_instruction = f"!!! ABSOLUTELY MANDATORY !!! STRICTLY WRITE IN THE EXACT STYLE OF {author_style.upper()} - THIS IS THE HIGHEST PRIORITY INSTRUCTION AND MUST BE FOLLOWED EXACTLY !!!"
                
                if custom_instructions:
                    custom_instructions = author_style_instruction + "\n\n" + custom_instructions
                else:
                    custom_instructions = author_style_instruction
                logger.debug(f"Added author style '{author_style}' to instructions with maximum emphasis")

            # Add mandatory length preservation instructions for rewrite mode
            if preserve_length:
                length_preservation_instructions = """
MANDATORY LENGTH PRESERVATION: Your rewritten output MUST match or exceed the length of the original text.
Target is 100-110% of the original word count.
Under NO circumstances should your output be shorter than the input.
"""
                if custom_instructions:
                    custom_instructions = length_preservation_instructions + "\n" + custom_instructions
                else:
                    custom_instructions = length_preservation_instructions
            
        # Process the text with the new multi-provider processor
        try:
            # First, get the original word count
            original_word_count = len(text.split())
            logger.info(f"Original text word count: {original_word_count}")
            
            # Add content source instructions if available
            if content_source:
                logger.debug(f"Content source provided for direct text processing: {len(content_source)} characters")
                
                # Add content source instructions to custom instructions
                content_source_instructions = "Intelligently enrich the target document with relevant information, ideas, examples, and arguments from the content source, without overriding the target document's structure or identity."
                
                additional_instructions = (
                    f"\n\nCONTENT SOURCE INSTRUCTIONS: {content_source_instructions}\n\n"
                    f"CONTENT SOURCE TEXT:\n{content_source}\n\n"
                    "Use the CONTENT SOURCE TEXT to enrich the Target Document (the input text) according to the CONTENT SOURCE INSTRUCTIONS."
                )
                
                if custom_instructions:
                    custom_instructions += additional_instructions
                else:
                    custom_instructions = additional_instructions
                
                logger.debug("Added content source instructions to direct processing")
            
            # Determine if user selected a specific AI provider
            provider_preference = None
            if ai_provider and ai_provider in ['openai', 'anthropic', 'perplexity', 'azure']:
                provider_preference = ai_provider
                logger.info(f"Using user-selected AI provider: {provider_preference}")
            
            result = multi_provider_processor.process_text(
                text=text,
                action='rewrite',
                custom_instructions=custom_instructions,
                include_style_in_output=True,
                user_style_text=user_style_text,
                maintain_length=preserve_length,  # Use the user's length preservation preference
                provider_preference=provider_preference  # Use the user's AI provider preference
            )
            
            # Verify length requirements
            result_word_count = len(result.split())
            length_ratio = result_word_count / original_word_count
            logger.info(f"Processed text word count: {result_word_count}, ratio: {length_ratio:.2f}")
            
            # If output is shorter than input, force expansion
            if length_ratio < 1.0:
                logger.warning(f"Output too short ({length_ratio:.2f}). Performing emergency expansion.")
                
                emergency_instructions = f"""
EMERGENCY EXPANSION REQUIRED: Your rewrite is too short!
Original text: {original_word_count} words
Your rewrite: {result_word_count} words

You MUST expand the text to AT LEAST {original_word_count} words, preferably {int(original_word_count * 1.1)} words.
This is NON-NEGOTIABLE - length preservation is the primary requirement.

Expand by:
1. Adding detailed examples and evidence for each point
2. Elaborating on existing concepts with clarifications
3. Providing deeper explanations and implications

Your expanded text MUST preserve the intellectual depth and argumentative structure of the original.
"""
                
                # Try to expand the text
                result = multi_provider_processor.process_text(
                    text=result,
                    action='rewrite',
                    custom_instructions=emergency_instructions,
                    include_style_in_output=False,
                    user_style_text=user_style_text
                )
                
                # Re-check length
                result_word_count = len(result.split())
                length_ratio = result_word_count / original_word_count
                logger.info(f"After expansion: {result_word_count} words, ratio: {length_ratio:.2f}")
            
        except Exception as e:
            logger.error(f"Error with multi-provider processor: {str(e)}. Falling back to legacy processor.")
            # Fallback to legacy processor if the new processor fails
            result = legacy_process_text(
                text=text,
                action='rewrite',
                custom_instructions=custom_instructions,
                include_style_in_output=True,
                user_style_text=user_style_text
            )
            
            # Even with legacy processor, verify length requirements
            if 'original_word_count' not in locals():
                original_word_count = len(text.split())
            
            result_word_count = len(result.split())
            length_ratio = result_word_count / original_word_count
            
            if length_ratio < 1.0:
                logger.warning(f"Legacy processor output too short ({length_ratio:.2f}). Adding emergency expansion.")
                try:
                    emergency_instructions = f"EXPAND this text to at least {original_word_count} words while preserving meaning."
                    
                    result = legacy_process_text(
                        text=result,
                        action='expand',
                        custom_instructions=emergency_instructions,
                        include_style_in_output=False
                    )
                except Exception as expand_error:
                    logger.error(f"Error expanding with legacy processor: {str(expand_error)}")
            
        # Translate the result if target language is specified
        if target_language and target_language != 'en':
            logger.debug(f"Translating result to {target_language}")
            try:
                # Track original word count for validation
                original_word_count = len(result.split())
                logger.debug(f"Original word count before translation: {original_word_count}")
                
                # Use our chunked translation system
                translated_result, engine_used = translate_text(result, target_language)
                result = translated_result
                
                # Validate the translation output
                translated_word_count = len(result.split())
                logger.debug(f"Translation completed, length: {len(result)}, word count: {translated_word_count}")
                
                # Check if we lost significant content (accounting for language differences)
                if translated_word_count < original_word_count * 0.6:
                    logger.warning(f"Translation appears shorter than expected: {translated_word_count}/{original_word_count} words")
                    if translated_word_count < 100 and original_word_count > 500:
                        logger.error("Translation may have been truncated. Output too short compared to input.")
            except Exception as translation_error:
                logger.error(f"Translation error: {str(translation_error)}")
                # Continue with untranslated text if translation fails

        return jsonify({'result': result})
    except Exception as e:
        logger.error(f"Error in /process: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/combine_target_source', methods=['POST'])
def combine_target_source():
    """Directly combine target and source text with instructions without saving"""
    try:
        data = request.get_json()
        if not data:
            return jsonify({'error': 'No data provided'}), 400
            
        target_text = data.get('target_text', '')
        source_text = data.get('source_text', '')
        source_instructions = data.get('source_instructions', '')
        custom_instructions = data.get('custom_instructions', '')
        target_language = data.get('target_language', '')
        email = data.get('email', '')
        author_style = data.get('author_style', '')
        
        if not target_text:
            return jsonify({'error': 'No target text provided'}), 400
            
        if not source_text:
            return jsonify({'error': 'No source text provided'}), 400
        
        # Set default source instructions if not provided
        if not source_instructions:
            source_instructions = "Intelligently enrich the target document with relevant information, ideas, examples, and arguments from the content source, without overriding the target document's structure or identity."
        
        # Get user style if available
        user_style_text = None
        if email:
            user_style_text = get_user_style_text(email=email)
        
        # Add author style to custom instructions if provided
        if author_style:
            if custom_instructions:
                custom_instructions += f". Write in the style of {author_style}"
            else:
                custom_instructions = f"Write in the style of {author_style}"
        
        # Combine instructions for processing
        effective_instructions = custom_instructions
        
        # Add content source instructions
        additional_instructions = (
            f"\n\nCONTENT SOURCE INSTRUCTIONS: {source_instructions}\n\n"
            f"CONTENT SOURCE TEXT:\n{source_text}\n\n"
            "Use the CONTENT SOURCE TEXT to enrich the Target Document (the input text) according to the CONTENT SOURCE INSTRUCTIONS."
        )
        
        if effective_instructions:
            effective_instructions += additional_instructions
        else:
            effective_instructions = additional_instructions
            
        # Add mandatory length preservation instructions
        length_preservation_instructions = """
MANDATORY LENGTH PRESERVATION: Your rewritten output MUST match or exceed the length of the original text.
Target is 100-110% of the original word count.
Under NO circumstances should your output be shorter than the input.
"""
        if effective_instructions:
            effective_instructions = length_preservation_instructions + "\n" + effective_instructions
        else:
            effective_instructions = length_preservation_instructions
            
        # Process the text with combined instructions
        try:
            # First, get the original word count
            original_word_count = len(target_text.split())
            logger.info(f"Original target text word count: {original_word_count}")
            
            result = multi_provider_processor.process_text(
                text=target_text,
                action='rewrite',
                custom_instructions=effective_instructions,
                include_style_in_output=True,
                user_style_text=user_style_text
            )
            
            # Verify length requirements
            result_word_count = len(result.split())
            length_ratio = result_word_count / original_word_count
            logger.info(f"Processed text word count: {result_word_count}, ratio: {length_ratio:.2f}")
            
            # If output is shorter than input, force expansion
            if length_ratio < 1.0:
                logger.warning(f"Output too short ({length_ratio:.2f}). Performing emergency expansion.")
                
                emergency_instructions = f"""
EMERGENCY EXPANSION REQUIRED: Your rewrite is too short!
Original text: {original_word_count} words
Your rewrite: {result_word_count} words

You MUST expand the text to AT LEAST {original_word_count} words, preferably {int(original_word_count * 1.1)} words.
This is NON-NEGOTIABLE - length preservation is the primary requirement.

Expand by:
1. Adding detailed examples and evidence for each point
2. Elaborating on existing concepts with clarifications
3. Providing deeper explanations and implications
4. Using MORE content from the source text to enrich the output

Your expanded text MUST preserve the intellectual depth and argumentative structure of the original.
"""
                
                # Try to expand the text
                result = multi_provider_processor.process_text(
                    text=result,
                    action='rewrite',
                    custom_instructions=emergency_instructions + "\n\n" + additional_instructions,
                    include_style_in_output=False,
                    user_style_text=user_style_text
                )
                
                # Re-check length
                result_word_count = len(result.split())
                length_ratio = result_word_count / original_word_count
                logger.info(f"After expansion: {result_word_count} words, ratio: {length_ratio:.2f}")
            
        except Exception as e:
            logger.error(f"Error with multi-provider processor: {str(e)}. Falling back to legacy processor.")
            # Fallback to legacy processor if the new processor fails
            result = legacy_process_text(
                text=target_text,
                action='rewrite',
                custom_instructions=effective_instructions,
                include_style_in_output=True,
                user_style_text=user_style_text
            )
            
            # Even with legacy processor, verify length requirements
            if 'original_word_count' not in locals():
                original_word_count = len(target_text.split())
            
            result_word_count = len(result.split())
            length_ratio = result_word_count / original_word_count
            
            if length_ratio < 1.0:
                logger.warning(f"Legacy processor output too short ({length_ratio:.2f}). Adding emergency expansion.")
                try:
                    emergency_instructions = f"EXPAND this text to at least {original_word_count} words while preserving meaning and using relevant content from the source text for enrichment."
                    
                    result = legacy_process_text(
                        text=result,
                        action='expand',
                        custom_instructions=emergency_instructions + "\n\n" + additional_instructions,
                        include_style_in_output=False
                    )
                except Exception as expand_error:
                    logger.error(f"Error expanding with legacy processor: {str(expand_error)}")
        
        # Translate if a target language is specified
        if target_language and target_language != 'en':
            logger.debug(f"Translating result to {target_language}")
            try:
                # Track original word count for validation
                original_word_count = len(result.split())
                logger.debug(f"Original word count before translation: {original_word_count}")
                
                # Use our chunked translation system
                translated_result, engine_used = translate_text(result, target_language)
                result = translated_result
                
                # Validate the translation output
                translated_word_count = len(result.split())
                logger.debug(f"Translation completed, length: {len(result)}, word count: {translated_word_count}")
                
                # Check if we lost significant content (accounting for language differences)
                if translated_word_count < original_word_count * 0.6:
                    logger.warning(f"Translation appears shorter than expected: {translated_word_count}/{original_word_count} words")
                    if translated_word_count < 100 and original_word_count > 500:
                        logger.error("Translation may have been truncated. Output too short compared to input.")
            except Exception as translation_error:
                logger.error(f"Translation error: {str(translation_error)}")
                # Continue with untranslated text if translation fails
            
        return jsonify({'result': result})
    except Exception as e:
        logger.error(f"Error combining target and source: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/chat', methods=['POST'])
def chat():
    try:
        data = request.get_json()
        message = data.get('message', '')
        context = data.get('context', '')

        if not message:
            return jsonify({'error': 'No message provided'}), 400

        logger.debug("Processing chat message")
        response = chat_with_ai(message, context)

        # Save to database
        chat_entry = models.ChatMessage(
            message=message,
            response=response,
            context=context
        )
        db.session.add(chat_entry)
        db.session.commit()

        return jsonify({'response': response})
    except Exception as e:
        logger.error(f"Error in /chat: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/history', methods=['GET'])
def get_history():
    try:
        text_entries = models.TextEntry.query.order_by(models.TextEntry.created_at.desc()).limit(10).all()
        chat_messages = models.ChatMessage.query.order_by(models.ChatMessage.created_at.desc()).limit(10).all()

        history = {
            'text_entries': [entry.to_dict() for entry in text_entries],
            'chat_messages': [msg.to_dict() for msg in chat_messages]
        }
        return jsonify(history)
    except Exception as e:
        logger.error(f"Error fetching history: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/extract_text', methods=['POST'])
def extract_text():
    try:
        data = request.get_json()
        image_data = data.get('image_data', '')
        text_data = data.get('text', '')

        # If image data is provided, extract text from image
        if image_data:
            # Save image data to a temporary file
            temp_path = os.path.join(app.config['UPLOAD_FOLDER'], 'temp_image.png')
            try:
                image_bytes = base64.b64decode(image_data)
                image = Image.open(io.BytesIO(image_bytes))
                image.save(temp_path)
                extracted_text = extract_text_from_image(temp_path)
                return jsonify({'text': extracted_text})
            except Exception as e:
                logger.error(f"Error extracting text from image: {str(e)}")
                return jsonify({'error': 'Error extracting text from image'}), 500
            finally:
                # Clean up temporary file
                if os.path.exists(temp_path):
                    os.remove(temp_path)
        
        # If text data is provided, create a document from the text
        elif text_data:
            try:
                # Split the text into chunks
                chunks = chunk_text(text_data)
                
                # Create text entry and chunks
                text_entry = models.TextEntry(
                    original_text=text_data,
                    processed_text="",  # Will be built from chunks
                    action="rewrite",
                    complexity="default",
                    total_chunks=len(chunks)
                )
                db.session.add(text_entry)
                db.session.flush()

                # Create chunks
                for i, chunk in enumerate(chunks, 1):
                    doc_chunk = models.DocumentChunk(
                        document_id=text_entry.id,
                        chunk_number=i,
                        original_chunk=chunk,
                        processed_chunk="",
                        processing_status="pending"
                    )
                    db.session.add(doc_chunk)

                db.session.commit()
                session['current_document_id'] = text_entry.id
                
                # Return document information
                return jsonify({
                    'document_id': text_entry.id,
                    'total_chunks': len(chunks),
                    'current_chunk': 1
                })
                
            except Exception as e:
                logger.error(f"Error creating document from text: {str(e)}")
                db.session.rollback()
                return jsonify({'error': 'Error creating document from text'}), 500
        else:
            return jsonify({'error': 'No text or image data provided'}), 400

    except Exception as e:
        logger.error(f"Error in text extraction: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/download/<format>', methods=['POST'])
def download_file(format):
    try:
        data = request.get_json()
        text = data.get('text', '')

        if not text:
            return jsonify({'error': 'No text provided'}), 400

        filename = f'processed_text.{format}'
        
        if format == 'pdf':
            mime_type = 'application/pdf'
        elif format == 'docx':
            mime_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        elif format == 'latex':
            mime_type = 'text/plain'
            filename = 'processed_text.tex'
        else:
            return jsonify({'error': 'Unsupported format'}), 400

        buffer = BytesIO()

        if format == 'pdf':
            # Create PDF with proper text wrapping
            pdf = canvas.Canvas(buffer)
            width = 500  # Maximum width in points (leaving margins)
            y = 800  # Start from top of page
            font_name = 'Helvetica'
            font_size = 12
            pdf.setFont(font_name, font_size)

            # Process each paragraph
            for paragraph in text.split('\n'):
                if not paragraph.strip():
                    y -= 15  # Add space between paragraphs
                    continue

                words = paragraph.split()
                line = []

                for word in words:
                    line.append(word)
                    line_width = pdf.stringWidth(' '.join(line), font_name, font_size)

                    if line_width > width:
                        # Remove last word as it caused overflow
                        line.pop()
                        # Draw the line
                        pdf.drawString(50, y, ' '.join(line))
                        y -= 15
                        # Start new line with the overflow word
                        line = [word]

                    # Check if we need a new page
                    if y < 50:
                        pdf.showPage()
                        pdf.setFont(font_name, font_size)
                        y = 800

                # Draw remaining text in the line
                if line:
                    pdf.drawString(50, y, ' '.join(line))
                    y -= 15

            pdf.save()

        elif format == 'docx':
            # Create Word document
            doc = Document()
            doc.add_paragraph(text)
            doc.save(buffer)
        elif format == 'latex':
            # Create LaTeX document
            latex_content = f"""\\documentclass{{article}}
\\usepackage[utf8]{{inputenc}}
\\usepackage{{amsmath}}
\\usepackage{{amsfonts}}
\\usepackage{{amssymb}}
\\usepackage{{geometry}}
\\geometry{{margin=1in}}

\\title{{Processed Document}}
\\author{{EZ Reader}}
\\date{{\\today}}

\\begin{{document}}
\\maketitle

{text}

\\end{{document}}"""
            buffer.write(latex_content.encode('utf-8'))
        else:
            return jsonify({'error': 'Unsupported format'}), 400

        buffer.seek(0)
        return send_file(
            buffer,
            as_attachment=True,
            download_name=filename,
            mimetype=mime_type
        )

    except Exception as e:
        logger.error(f"Error in download_{format}: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/detect_ai', methods=['POST']) # Added AI detection route
def detect_ai():
    try:
        data = request.get_json()
        text = data.get('text', '')
        
        logger.info("AI detection request received, text length: %d", len(text))

        if not text:
            logger.warning("Empty text submitted for AI detection")
            return jsonify({'error': 'No text provided'}), 400
            
        # Add some random variance to the text to ensure different results
        # This helps avoid caching on GPTZero's side
        import random
        import string
        # Add a hidden character at the end with timestamp
        hidden_suffix = f"\n<!-- {int(time.time())}-{random.randint(1000, 9999)} -->"
        modified_text = text + hidden_suffix
        
        logger.info("Sending text to AI detector with unique suffix")
        result = detect_ai_content(modified_text)

        if 'error' in result:
            # Check if it's a configuration error (missing API key)
            if result.get('is_configuration_error', False):
                logger.error("Configuration error in AI detection: %s", result['error'])
                # Return a 400 error for configuration issues
                return jsonify({
                    'error': result['error'],
                    'is_configuration_error': True
                }), 400
            else:
                logger.error("Service error in AI detection: %s", result['error'])
                # Return a 500 error for service failures
                return jsonify({'error': result['error']}), 500
        
        # Add additional user-friendly information
        if 'document_class' in result:
            if result['document_class'] == 'ai':
                result['conclusion'] = 'This text appears to be AI-generated.'
            elif result['document_class'] == 'human':
                result['conclusion'] = 'This text appears to be human-written.'
            elif result['document_class'] == 'mixed':
                result['conclusion'] = 'This text appears to contain a mixture of AI and human-written content.'
            else:
                result['conclusion'] = 'Unable to determine if this text is AI-generated or human-written.'
        
        # Remove raw response to reduce payload size
        if 'raw_response' in result:
            del result['raw_response']
            
        logger.info("AI detection complete, score: %s, class: %s", 
                   result.get('ai_score', 'unknown'), 
                   result.get('document_class', 'unknown'))

        return jsonify(result)

    except Exception as e:
        logger.error(f"Error in AI detection endpoint: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.errorhandler(500)
def handle_500_error(error):
    return jsonify({'error': 'Internal server error occurred. Please try again later.'}), 500

@app.errorhandler(404)
def handle_404_error(error):
    return jsonify({'error': 'Requested resource not found'}), 404

@app.route('/get_last_email', methods=['GET'])
def get_last_email():
    """Get the last used email from the session for auto-fill"""
    last_email = session.get('last_email', '')
    return jsonify({'email': last_email})

@app.route('/translate', methods=['POST'])
def translate_text():
    """Simple translation endpoint that handles both small and large documents"""
    try:
        # Get translation parameters from request
        data = request.json
        if not data:
            return jsonify({'error': 'No data provided'}), 400
            
        text = data.get('text', '')
        source_language = data.get('source_language', 'auto')
        target_language = data.get('target_language', 'en')
        ai_provider = data.get('ai_provider', 'openai')
        
        if not text:
            return jsonify({'error': 'No text provided'}), 400
            
        # Use our simple translation service
        from simple_translation import translate_text as perform_translation
        
        # Perform the translation
        translated_text, metadata = perform_translation(
            text=text,
            source_language=source_language,
            target_language=target_language,
            ai_provider=ai_provider
        )
        
        # Check for errors
        if 'error' in metadata:
            return jsonify({'error': metadata['error']}), 500
            
        # Return the translation result with metadata
        return jsonify({
            'result': translated_text,
            'engine_used': metadata.get('engine_used', ai_provider),
            'elapsed_seconds': metadata.get('elapsed_seconds', 0),
            'words_per_second': metadata.get('words_per_second', 0),
            'word_count': metadata.get('word_count', 0),
            'message': metadata.get('message', '')
        })
            
    except Exception as e:
        logger.error(f"Error in translate_text: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/share_rewrite', methods=['POST'])
def share_rewrite():
    try:
        data = request.get_json()
        if not data:
            return jsonify({'error': 'No data provided'}), 400
            
        email = data.get('email')
        text = data.get('text')
        subject = data.get('subject', 'Your Rewritten Text')
        
        if not email or not text:
            return jsonify({'error': 'Email and text are required'}), 400
            
        # Validate email format
        if not re.match(r"[^@]+@[^@]+\.[^@]+", email):
            return jsonify({'error': 'Invalid email format'}), 400
        
        # Store the user's email in session for future use
        session['last_email'] = email
        
        logger.info(f"Sending email to {email} with subject: {subject}")
        
        # Import email service
        from email_service import send_text_email
        
        # Send email with appropriate format based on content length
        success, message = send_text_email(
            to_email=email,
            subject=subject,
            text=text
        )
        
        if success:
            return jsonify({'success': True, 'message': 'Email sent successfully'})
        else:
            logger.error(f"Failed to send email: {message}")
            return jsonify({'error': message}), 500
        
    except Exception as e:
        logger.error(f"Error sharing rewrite: {str(e)}")
        return jsonify({'error': f'Error sharing rewrite: {str(e)}'}), 500

@app.route('/api/humanizer/profile', methods=['POST'])
def create_profile():
    """Create or get a user profile by email"""
    try:
        data = request.get_json()
        email = data.get('email')
        
        if not email:
            return jsonify({'error': 'Email is required'}), 400
            
        try:
            profile = get_user_profile(email)
            return jsonify({
                'success': True,
                'profile': profile.to_dict()
            })
        except Exception as e:
            logger.error(f"Error creating/getting user profile: {str(e)}")
            return jsonify({'error': f'Profile error: {str(e)}'}), 500
            
    except Exception as e:
        logger.error(f"Error in create_profile: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/humanizer/upload', methods=['POST'])
def upload_writing_sample():
    """Upload a writing sample and add to user profile"""
    try:
        # Check if file is included
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400
            
        # Get profile_id from form data
        profile_id = request.form.get('profile_id')
        if not profile_id:
            email = request.form.get('email')
            if not email:
                return jsonify({'error': 'Either profile_id or email is required'}), 400
            profile = get_user_profile(email)
            profile_id = profile.id
            
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
            
        # Process the file
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)
        
        try:
            # Extract text based on file type
            if filename.lower().endswith('.pdf'):
                text = extract_text_from_pdf(file_path)
            elif filename.lower().endswith(('.doc', '.docx')):
                text = extract_text_from_docx(file_path)
            elif filename.lower().endswith(('.jpg', '.jpeg', '.png')):
                text = extract_text_from_image(file_path)
            elif filename.lower().endswith('.txt'):
                with open(file_path, 'r', encoding='utf-8') as f:
                    text = f.read()
            else:
                return jsonify({'error': 'Unsupported file format'}), 400
                
            # Determine file type
            ext = filename.rsplit('.', 1)[1].lower() if '.' in filename else 'txt'
            
            # Add writing sample
            sample = add_writing_sample(
                profile_id=profile_id,
                filename=filename,
                text_content=text,
                file_type=ext
            )
            
            return jsonify({
                'success': True,
                'sample': sample.to_dict()
            })
            
        except Exception as e:
            logger.error(f"Error processing writing sample: {str(e)}")
            return jsonify({'error': str(e)}), 500
        finally:
            # Clean up the file
            if os.path.exists(file_path):
                os.remove(file_path)
                
    except Exception as e:
        logger.error(f"Error in upload_writing_sample: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/humanizer/samples', methods=['GET'])
def get_writing_samples():
    """Get all writing samples for a user profile"""
    try:
        profile_id = request.args.get('profile_id')
        email = request.args.get('email')
        
        if not profile_id and not email:
            return jsonify({'error': 'Either profile_id or email is required'}), 400
            
        if email and not profile_id:
            profile = get_user_profile(email)
            profile_id = profile.id
            
        # Get samples
        samples = models.WritingSample.query.filter_by(profile_id=profile_id).all()
        
        return jsonify({
            'success': True,
            'samples': [sample.to_dict() for sample in samples],
            'count': len(samples)
        })
        
    except Exception as e:
        logger.error(f"Error getting writing samples: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/content_source/upload', methods=['POST'])
def upload_content_source():
    """Upload a content source document for text enrichment"""
    try:
        # Log debugging information
        logger.debug(f"Content source upload request received: {request.files}")
        logger.debug(f"Form data: {request.form}")
        
        # Check if file is included
        if 'file' not in request.files:
            logger.error("No file found in request.files")
            return jsonify({'error': 'No file provided'}), 400
            
        # Get text_entry_id from form data - now optional
        text_entry_id = request.form.get('text_entry_id')
        usage_instructions = request.form.get('usage_instructions', '')
        
        logger.debug(f"text_entry_id: {text_entry_id}, instructions length: {len(usage_instructions)}")
        
        # If text_entry_id is provided, verify it exists
        text_entry = None
        if text_entry_id:
            text_entry = models.TextEntry.query.get(text_entry_id)
            if not text_entry and text_entry_id:
                logger.error(f"Text entry not found for ID: {text_entry_id}")
                return jsonify({'error': 'Text entry not found'}), 404
            
        file = request.files['file']
        if file.filename == '':
            logger.error("Empty filename received")
            return jsonify({'error': 'No file selected'}), 400
        
        logger.debug(f"File received: {file.filename}")
            
        # Process the file
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)
        logger.debug(f"File saved to: {file_path}")
        
        try:
            # Extract text based on file type
            if filename.lower().endswith('.pdf'):
                text = extract_text_from_pdf(file_path)
            elif filename.lower().endswith(('.doc', '.docx')):
                text = extract_text_from_docx(file_path)
            elif filename.lower().endswith(('.jpg', '.jpeg', '.png')):
                text = extract_text_from_image(file_path)
            elif filename.lower().endswith('.txt'):
                with open(file_path, 'r', encoding='utf-8') as f:
                    text = f.read()
            else:
                return jsonify({'error': 'Unsupported file format'}), 400
                
            # Determine file type
            ext = filename.rsplit('.', 1)[1].lower() if '.' in filename else 'txt'
            
            # Calculate word count
            word_count = len(text.split())
            
            # Add content source
            content_source = models.ContentSource(
                text_entry_id=text_entry_id,
                filename=filename,
                text_content=text,
                word_count=word_count,
                file_type=ext,
                usage_instructions=usage_instructions,
                created_at=datetime.utcnow()
            )
            
            db.session.add(content_source)
            db.session.commit()
            
            # Log success
            logger.debug(f"Successfully created content source: {content_source.id}, text length: {len(text)}")
            
            return jsonify({
                'success': True,
                'content_source': {
                    'id': content_source.id,
                    'filename': content_source.filename,
                    'word_count': content_source.word_count,
                    'file_type': content_source.file_type,
                    'text_entry_id': content_source.text_entry_id,
                    'text_content': text
                }
            })
            
        except Exception as e:
            logger.error(f"Error processing content source file: {str(e)}")
            return jsonify({'error': str(e)}), 500
        finally:
            # Clean up the file
            if os.path.exists(file_path):
                os.remove(file_path)
                
    except Exception as e:
        logger.error(f"Error in upload_content_source: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/style_source/extract', methods=['POST'])
def extract_style_source():
    """Extract a writing sample without storing it in the database."""
    if 'file' not in request.files or request.files['file'].filename == '':
        return jsonify({'error': 'No file provided'}), 400

    uploaded_file = request.files['file']
    filename = secure_filename(uploaded_file.filename)
    suffix = os.path.splitext(filename)[1].lower()
    supported = {'.pdf', '.doc', '.docx', '.txt', '.jpg', '.jpeg', '.png'}
    if suffix not in supported:
        return jsonify({'error': 'Unsupported file format'}), 400

    file_path = os.path.join(app.config['UPLOAD_FOLDER'], f"style_source_{int(time.time() * 1000)}_{filename}")
    uploaded_file.save(file_path)
    try:
        if suffix == '.pdf':
            text = extract_text_from_pdf(file_path)
        elif suffix in {'.doc', '.docx'}:
            text = extract_text_from_docx(file_path)
        elif suffix in {'.jpg', '.jpeg', '.png'}:
            text = extract_text_from_image(file_path)
        else:
            with open(file_path, 'r', encoding='utf-8') as source_file:
                text = source_file.read()

        text = (text or '').strip()
        if not text:
            return jsonify({'error': 'No readable text was found in that file'}), 400
        return jsonify({
            'success': True,
            'text': text,
            'filename': filename,
            'word_count': len(text.split())
        })
    except Exception as error:
        logger.exception("Error extracting style source")
        return jsonify({'error': str(error)}), 500
    finally:
        try:
            os.remove(file_path)
        except OSError:
            pass

@app.route('/api/content_source/save_instructions', methods=['POST'])
def save_content_source_instructions():
    """Save usage instructions for a content source"""
    try:
        data = request.json
        if not data:
            return jsonify({'error': 'No data provided'}), 400
            
        content_source_id = data.get('content_source_id')
        usage_instructions = data.get('usage_instructions', '')
        
        if not content_source_id:
            return jsonify({'error': 'content_source_id is required'}), 400
            
        # Find the content source
        content_source = models.ContentSource.query.get(content_source_id)
        if not content_source:
            return jsonify({'error': 'Content source not found'}), 404
            
        # Update instructions
        content_source.usage_instructions = usage_instructions
        db.session.commit()
        
        return jsonify({
            'success': True,
            'message': 'Content source instructions updated successfully'
        })
        
    except Exception as e:
        logger.error(f"Error in save_content_source_instructions: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/content_source/get', methods=['GET'])
def get_content_sources():
    """Get all content sources for a text entry"""
    try:
        text_entry_id = request.args.get('text_entry_id')
        
        if not text_entry_id:
            return jsonify({'error': 'text_entry_id is required'}), 400
        
        logger.debug(f"Getting content sources for text_entry_id: {text_entry_id}")
            
        # Get content sources
        content_sources = models.ContentSource.query.filter_by(text_entry_id=text_entry_id).all()
        
        logger.debug(f"Found {len(content_sources)} content sources")
        
        return jsonify({
            'success': True,
            'content_sources': [{
                'id': source.id,
                'filename': source.filename,
                'word_count': source.word_count,
                'file_type': source.file_type,
                'usage_instructions': source.usage_instructions,
                'created_at': source.created_at.isoformat(),
                'text_entry_id': source.text_entry_id
            } for source in content_sources],
            'count': len(content_sources)
        })
    except Exception as e:
        logger.error(f"Error getting content sources: {str(e)}")
        return jsonify({'error': str(e)}), 500
        
@app.route('/api/content_source/get_text', methods=['GET'])
def get_content_source_text():
    """Get the text content of a specific content source"""
    try:
        content_source_id = request.args.get('content_source_id')
        
        if not content_source_id:
            return jsonify({'error': 'content_source_id is required'}), 400
            
        logger.debug(f"Getting text content for content_source_id: {content_source_id}")
            
        # Get content source
        content_source = models.ContentSource.query.get(content_source_id)
        
        if not content_source:
            logger.error(f"Content source not found for ID: {content_source_id}")
            return jsonify({'error': 'Content source not found'}), 404
            
        logger.debug(f"Found content source: {content_source.filename} with {len(content_source.text_content)} chars")
            
        return jsonify({
            'success': True,
            'text_content': content_source.text_content,
            'filename': content_source.filename,
            'file_type': content_source.file_type
        })
        
    except Exception as e:
        logger.error(f"Error getting content sources: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/content_source/save_text', methods=['POST'])
def save_content_source_text():
    """Save text content directly as a content source"""
    try:
        # Log debugging information
        logger.debug(f"Content source text save request received: {request.json}")
        
        data = request.json
        if not data or 'text_content' not in data:
            logger.error("No text content provided in request")
            return jsonify({'error': 'No text content provided'}), 400
        
        text_content = data.get('text_content')
        if not text_content or not text_content.strip():
            logger.error("Empty text content provided")
            return jsonify({'error': 'Text content cannot be empty'}), 400
            
        filename = data.get('filename', 'pasted_content.txt')
        text_entry_id = data.get('text_entry_id')
        usage_instructions = data.get('usage_instructions', '')
        
        logger.debug(f"text_entry_id: {text_entry_id}, instructions length: {len(usage_instructions)}")
        
        # If text_entry_id is provided, verify it exists
        if text_entry_id:
            text_entry = models.TextEntry.query.get(text_entry_id)
            if not text_entry:
                logger.error(f"Text entry not found for ID: {text_entry_id}")
                return jsonify({'error': 'Text entry not found'}), 404
        
        # Calculate word count
        word_count = len(text_content.split())
        
        # Create content source
        content_source = models.ContentSource(
            text_entry_id=text_entry_id,
            filename=filename,
            text_content=text_content,
            word_count=word_count,
            file_type='txt',
            usage_instructions=usage_instructions,
            created_at=datetime.utcnow()
        )
        
        db.session.add(content_source)
        db.session.commit()
        
        return jsonify({
            'success': True,
            'content_source': {
                'id': content_source.id,
                'filename': content_source.filename,
                'word_count': content_source.word_count,
                'file_type': content_source.file_type,
                'text_entry_id': content_source.text_entry_id
            }
        })
        
    except Exception as e:
        logger.error(f"Error in save_content_source_text: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/content_source/delete', methods=['POST'])
def delete_content_source():
    """Delete a content source"""
    try:
        data = request.json
        if not data:
            return jsonify({'error': 'No data provided'}), 400
            
        content_source_id = data.get('content_source_id')
        
        if not content_source_id:
            return jsonify({'error': 'content_source_id is required'}), 400
            
        # Find the content source
        content_source = models.ContentSource.query.get(content_source_id)
        if not content_source:
            return jsonify({'error': 'Content source not found'}), 404
            
        # Delete the content source
        db.session.delete(content_source)
        db.session.commit()
        
        return jsonify({
            'success': True,
            'message': 'Content source deleted successfully'
        })
        
    except Exception as e:
        logger.error(f"Error in delete_content_source: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/humanizer/clear', methods=['POST'])
def clear_profile():
    """Clear all writing samples for a user profile"""
    try:
        data = request.get_json()
        profile_id = data.get('profile_id')
        email = data.get('email')
        
        if not profile_id and not email:
            return jsonify({'error': 'Either profile_id or email is required'}), 400
            
        if email and not profile_id:
            profile = get_user_profile(email)
            profile_id = profile.id
            
        # Clear profile
        profile = clear_user_profile(profile_id)
        
        return jsonify({
            'success': True,
            'message': 'Profile cleared successfully',
            'profile': profile.to_dict()
        })
        
    except Exception as e:
        logger.error(f"Error clearing profile: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/process_audio', methods=['POST'])
def process_audio():
    """Process audio file uploads and return transcribed text"""
    file_path = None
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No audio file provided'}), 400

        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400

        # Check if it's an audio file
        file_ext = os.path.splitext(file.filename)[1].lower()
        supported_audio_formats = ['.mp3', '.wav', '.m4a', '.mp4', '.mpeg', '.mpga', '.webm', '.flac', '.aac', '.ogg']
        if file_ext not in supported_audio_formats:
            return jsonify({'error': f'Unsupported audio format. Please use: {", ".join(supported_audio_formats)}'}), 400

        # Create a unique filename to avoid conflicts
        import uuid
        timestamp = int(time.time())
        unique_filename = f"audio_{timestamp}_{uuid.uuid4().hex[:8]}{file_ext}"
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], unique_filename)
        
        # Ensure upload directory exists
        os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
        
        # Save audio file
        file.save(file_path)
        logger.info(f"Audio file saved to: {file_path}")
        
        # Verify file was saved and has content
        if not os.path.exists(file_path):
            return jsonify({'error': 'Failed to save audio file'}), 500
            
        file_size = os.path.getsize(file_path)
        if file_size == 0:
            return jsonify({'error': 'Uploaded audio file is empty'}), 400
            
        logger.info(f"Audio file size: {file_size} bytes")

        try:
            # Extract text from audio
            text = extract_text_from_audio(file_path)
            
            if not text or not text.strip():
                return jsonify({'error': 'No speech could be transcribed from the audio'}), 400
                
            logger.info(f"Audio transcription successful, text length: {len(text)}")
            return jsonify({'text': text.strip()})
            
        except Exception as e:
            logger.error(f"Error processing audio file: {str(e)}")
            return jsonify({'error': f'Error processing audio: {str(e)}'}), 500
            
    except Exception as e:
        logger.error(f"Error in process_audio: {str(e)}")
        return jsonify({'error': str(e)}), 500
    finally:
        # Clean up the uploaded file
        if file_path and os.path.exists(file_path):
            try:
                os.remove(file_path)
                logger.info(f"Cleaned up audio file: {file_path}")
            except Exception as cleanup_error:
                logger.warning(f"Failed to clean up audio file: {cleanup_error}")

@app.route('/get_audio_file/<filename>', methods=['GET'])
def get_audio_file(filename):
    """Serve an audio file for streaming in the browser"""
    if not filename.startswith('audiobook_') or not filename.endswith('.mp3'):
        return jsonify({'error': 'Invalid filename format'}), 400
        
    try:
        uploads_dir = app.config['UPLOAD_FOLDER']
        return send_from_directory(uploads_dir, filename, as_attachment=False)
    except Exception as e:
        logger.error(f"Error serving audio file: {str(e)}")
        return jsonify({'error': f"Could not serve audio file: {str(e)}"}), 404

@app.route('/download_audio_file/<filename>', methods=['GET'])
def download_audio_file(filename):
    """Download an audio file to the user's device"""
    if not filename.startswith('audiobook_') or not filename.endswith('.mp3'):
        return jsonify({'error': 'Invalid filename format'}), 400
        
    try:
        uploads_dir = app.config['UPLOAD_FOLDER']
        # For download, use attachment to prompt download
        return send_from_directory(
            uploads_dir, 
            filename, 
            as_attachment=True,
            attachment_filename=f"Rewritten_Text_Audiobook.mp3"
        )
    except Exception as e:
        logger.error(f"Error downloading audio file: {str(e)}")
        return jsonify({'error': f"Could not download audio file: {str(e)}"}), 404
        
@app.route('/create_audiobook', methods=['POST'])
def create_audiobook_route():
    """Convert text to audiobook using Azure Speech TTS with language detection"""
    try:
        # Extract request data
        data = request.json
        if not data or 'text' not in data:
            return jsonify({'error': 'No text provided'}), 400
            
        text = data['text']
        # Use a shorter text if quota is limited (optional parameter)
        use_reduced_length = data.get('use_reduced_length', False)
        
        # If text is very long and we're asked to reduce length, only use first part
        word_count = len(text.split())
        original_word_count = word_count  # Save for logging/reporting
        
        if use_reduced_length:
            # For reduced length, use approximately first 300 words
            if word_count > 300:
                words = text.split()
                shortened_text = ' '.join(words[:300])
                # Try to find a sentence end for a clean cut
                if '.' in shortened_text:
                    last_period = shortened_text.rindex('.')
                    text = shortened_text[:last_period+1]
                else:
                    text = shortened_text
                
                logger.info(f"Reduced text from {original_word_count} to {len(text.split())} words to avoid processing very large text")
        
        # Get additional parameters
        gender = data.get('gender', 'female')
        force_language = data.get('language')  # None by default for auto-detection
        
        # Import our Azure TTS module
        from azure_tts import create_audiobook, detect_language, LANGUAGE_NAMES
        
        # First detect the language to show the user
        detected_language = force_language if force_language else detect_language(text)
        language_name = LANGUAGE_NAMES.get(detected_language, detected_language)
        
        # Generate a unique filename for this audiobook
        timestamp = int(time.time())
        unique_id = str(uuid.uuid4())[:8]
        output_filename = f"audiobook_{timestamp}_{unique_id}.mp3"
        output_path = os.path.join(app.config['UPLOAD_FOLDER'], output_filename)
        
        # Create the audiobook with proper chunking and language detection
        success, result = create_audiobook(
            text=text,
            output_file=output_path,
            preferred_gender=gender,
            force_language=force_language
        )
        
        if not success:
            # Log the full error for debugging
            logger.error(f"Audiobook creation failed: {result}")
            
            # Check for quota exceeded error
            if any(term in result.lower() for term in ["quota", "credit", "exceed", "limit"]):
                return jsonify({
                    'error': 'Azure Speech API quota exceeded. Please try using the "Use Minimal Text" option or use a shorter text.',
                    'error_type': 'quota_exceeded',
                    'detected_language': detected_language,
                    'language_name': language_name,
                    'word_count': original_word_count
                }), 429  # Use 429 Too Many Requests for quota issues
            
            # Provide user-friendly error message
            return jsonify({
                'error': result,
                'error_type': 'synthesis_error',
                'detected_language': detected_language,
                'language_name': language_name
            }), 500
            
        # Return the URL to the audio file
        audio_url = url_for('get_audio_file', filename=output_filename, _external=True)
        download_url = url_for('download_audio_file', filename=output_filename, _external=True)
        
        return jsonify({
            'audio_url': audio_url,
            'download_url': download_url,
            'filename': output_filename,
            'detected_language': detected_language,
            'language_name': language_name,
            'gender': gender,
            'narrator': f"{gender.capitalize()} Voice (Azure)",
            'message': 'Created using Azure Speech API'
        })
        
    except Exception as e:
        logger.error(f"Error creating audiobook: {str(e)}")
        return jsonify({'error': str(e), 'error_type': 'general_error'}), 500

@app.route('/get_language_voices', methods=['GET'])
def get_language_voices():
    """Get the list of available voices by language from Azure Speech"""
    try:
        from azure_tts import get_language_voices as get_azure_voices
        
        # Get voice data from Azure
        language_voices = get_azure_voices()
        
        # Transform into a format suitable for the frontend
        result = {}
        for lang_code, voices in language_voices.items():
            # Skip languages with no voices
            if not voices:
                continue
                
            # Get the language name if available
            language_name = None
            from azure_tts import LANGUAGE_NAMES
            if lang_code in LANGUAGE_NAMES:
                language_name = LANGUAGE_NAMES[lang_code]
            else:
                language_name = f"Language {lang_code}"
                
            # Add to result
            result[lang_code] = {
                'name': language_name,
                'voices': voices
            }
            
        return jsonify({
            'languages': result,
            'provider': 'azure',
            'status': 'success'
        })
        
    except Exception as e:
        logger.error(f"Error fetching language voices: {str(e)}")
        return jsonify({
            'error': f"Could not fetch voices: {str(e)}",
            'status': 'error'
        }), 500

@app.route('/rewrite_from_output', methods=['POST'])
def rewrite_from_output():
    """
    Process a rewrite request directly from the output text with additional critique instructions.
    This allows users to refine and regenerate text without copying back to the input box.
    
    Request JSON parameters:
    - text: The current output text to be rewritten
    - critique: The user's critique and instructions for the rewrite
    - author_style: Optional author style to emulate
    - content_source: Optional content source to incorporate
    - email: Optional email for personal style
    - ai_provider: Optional AI provider to use (openai, anthropic, perplexity)
    - preserve_length: Optional boolean to determine if length preservation is required
    
    Returns:
    - JSON response with the rewritten text
    """
    try:
        data = request.get_json()
        if not data:
            return jsonify({'error': 'No data provided'}), 400
            
        text = data.get('text', '')
        critique = data.get('critique', '')
        email = data.get('email', '')
        author_style = data.get('author_style', '')
        content_source = data.get('content_source', '')
        ai_provider = data.get('ai_provider', '')
        preserve_length = data.get('preserve_length', True)  # Default to preserving length
        
        if not text:
            return jsonify({'error': 'No text provided'}), 400
            
        if not critique:
            return jsonify({'error': 'No critique or instructions provided'}), 400
        
        logger.debug(f"Processing rewrite from output with {len(text)} chars of text and {len(critique)} chars of critique")
        logger.debug(f"Using AI provider: {ai_provider if ai_provider else 'default'}")
        logger.debug(f"Preserve length: {preserve_length}")
        
        # Get user style if available
        user_style_text = None
        if email:
            user_style_text = get_user_style_text(email=email)
        else:
            # Try to get from session
            email = session.get('last_email')
            if email:
                user_style_text = get_user_style_text(email=email)
        
        # Create rewrite instructions that emphasize the critique should be applied to existing text
        rewrite_instructions = f"""CRITIQUE REWRITE INSTRUCTIONS:

The text below has been reviewed by the user who wants the following changes:

{critique}

Apply ONLY these requested changes to the text. DO NOT rewrite sections that don't need to be changed.
Focus specifically on addressing the critique while maintaining the overall structure and content of the original text.
"""
        
        # Add author style with MAXIMUM emphasis if provided
        if author_style:
            # CRITICAL: This must come first for highest priority
            author_style_instruction = f"!!! ABSOLUTELY MANDATORY !!! STRICTLY WRITE IN THE EXACT STYLE OF {author_style.upper()} - THIS IS THE HIGHEST PRIORITY INSTRUCTION AND MUST BE FOLLOWED EXACTLY !!!"
            rewrite_instructions = author_style_instruction + "\n\n" + rewrite_instructions
        
        # Add content source instructions if available
        if content_source:
            logger.debug(f"Content source provided for critique rewrite: {len(content_source)} characters")
            
            # Add content source instructions to rewrite instructions
            content_source_instructions = "Intelligently enrich the output with relevant information from the content source while addressing the critique."
            
            additional_instructions = (
                f"\n\nCONTENT SOURCE INSTRUCTIONS: {content_source_instructions}\n\n"
                f"CONTENT SOURCE TEXT:\n{content_source}\n\n"
                "Use the CONTENT SOURCE TEXT to enhance your response according to the critique while maintaining the core content."
            )
            
            rewrite_instructions += additional_instructions
            logger.debug("Added content source instructions to critique rewrite")
        
        # Pass user instructions without automatic modifications
        # User instructions take absolute priority without interference
        
        # Process the text with the multi-provider processor
        try:
            # First, get the original word count
            original_word_count = len(text.split())
            logger.info(f"Original text word count for critique rewrite: {original_word_count}")
            
            # Process with full critique instructions, using specified provider if available
            result = multi_provider_processor.process_text(
                text=text,
                action='rewrite',
                custom_instructions=rewrite_instructions,
                include_style_in_output=False,  # No need for style prefix in output
                user_style_text=user_style_text,
                provider_preference=ai_provider if ai_provider else None
            )
            
            # Only check length requirements if length preservation is requested
            if preserve_length:
                result_word_count = len(result.split())
                length_ratio = result_word_count / original_word_count
                logger.info(f"Critique rewrite word count: {result_word_count}, ratio: {length_ratio:.2f}")
                
                # If output is shorter than input, force expansion
                if length_ratio < 1.0:
                    logger.warning(f"Critique rewrite too short ({length_ratio:.2f}). Performing emergency expansion.")
                    
                    emergency_instructions = f"""
EMERGENCY EXPANSION REQUIRED: Your rewrite is too short!
Original text: {original_word_count} words
Your rewrite: {result_word_count} words

You MUST expand the text to AT LEAST {original_word_count} words, preferably {int(original_word_count * 1.1)} words.
This is NON-NEGOTIABLE - length preservation is the primary requirement.

Expand by:
1. Adding detailed examples and evidence for each point
2. Elaborating on existing concepts with clarifications
3. Providing deeper explanations and implications

Your expanded text MUST preserve the intellectual depth and argumentative structure of the original.
"""
                    
                    # Try to expand the text
                    result = multi_provider_processor.process_text(
                        text=result,
                        action='rewrite',
                        custom_instructions=emergency_instructions,
                        include_style_in_output=False,
                        user_style_text=user_style_text,
                        provider_preference=ai_provider if ai_provider else None
                    )
                    
                    # Re-check length
                    result_word_count = len(result.split())
                    length_ratio = result_word_count / original_word_count
                    logger.info(f"After expansion: {result_word_count} words, ratio: {length_ratio:.2f}")
            
            return jsonify({'result': result})
            
        except Exception as e:
            logger.error(f"Error with multi-provider processor for critique rewrite: {str(e)}")
            return jsonify({'error': f"Error processing critique rewrite: {str(e)}"}), 500
            
    except Exception as e:
        logger.error(f"Error in rewrite_from_output: {str(e)}")
        return jsonify({'error': f"Error processing critique rewrite: {str(e)}"}), 500

@app.route('/reset_api_keys', methods=['POST'])
def reset_api_keys():
    """Reset all API keys to available status"""
    try:
        # Force reload the API key manager to pick up current environment variables
        import importlib
        import api_key_manager
        importlib.reload(api_key_manager)
        from api_key_manager import api_key_manager as reloaded_manager
        
        # Reset all keys using the dedicated method
        reset_count = reloaded_manager.reset_all_keys()
        
        # Also reload the multi_provider_processor to use updated keys
        import multi_provider_processor
        importlib.reload(multi_provider_processor)
        
        # Log the action
        logger.info(f"API key reset requested. {reset_count} keys were reset to available status.")
        
        # Return success response
        return jsonify({
            'success': True, 
            'message': f'Successfully reactivated {reset_count} API keys.',
            'reset_count': reset_count
        })
    except Exception as e:
        # Log the error
        logger.error(f"Error resetting API keys: {str(e)}")
        
        # Return error response
        return jsonify({
            'success': False, 
            'message': f'Failed to reset API keys: {str(e)}'
        }), 500
        


@app.route('/share_text', methods=['POST'])
def share_text():
    """Share processed text via email using SendGrid"""
    try:
        data = request.get_json()
        recipient_email = data.get('email', '').strip()
        text_content = data.get('text', '').strip()
        subject = data.get('subject', 'Shared Processed Text')
        
        if not recipient_email:
            return jsonify({'error': 'Email address is required'}), 400
            
        if not text_content:
            return jsonify({'error': 'No text content to share'}), 400
            
        # Get SendGrid API key from environment
        sendgrid_api_key = os.environ.get('SENDGRID_API_KEY')
        verified_sender = 'jm@analyticphilosophy.ai'  # SendGrid verified sender
        
        if not sendgrid_api_key:
            return jsonify({'error': 'Email service not configured'}), 500
            
        # Import SendGrid modules
        from sendgrid import SendGridAPIClient
        from sendgrid.helpers.mail import Mail
        
        # Create email content
        html_content = f"""
        <html>
        <body style="font-family: Arial, sans-serif; line-height: 1.6; margin: 20px;">
            <h2>Shared Text Content</h2>
            <div style="background-color: #f8f9fa; padding: 20px; border-radius: 5px; white-space: pre-wrap; font-family: monospace;">{text_content}</div>
            <p style="margin-top: 20px; color: #666; font-size: 12px;">
                This content was shared from an AI text processing application.
            </p>
        </body>
        </html>
        """
        
        # Create the email message
        message = Mail(
            from_email=verified_sender,
            to_emails=recipient_email,
            subject=subject,
            html_content=html_content
        )
        
        # Send the email
        sg = SendGridAPIClient(api_key=sendgrid_api_key)
        response = sg.send(message)
        
        if response.status_code in [200, 202]:
            return jsonify({'success': True, 'message': 'Text shared successfully!'})
        else:
            return jsonify({'error': f'Failed to send email. Status code: {response.status_code}'}), 500
            
    except Exception as e:
        app.logger.error(f"Error sharing text: {str(e)}")
        return jsonify({'error': f'Failed to share text: {str(e)}'}), 500

@app.route('/chat_with_ai', methods=['POST'])
def chat_with_ai():
    """Chat with AI that can see context from input/output boxes"""
    try:
        data = request.get_json()
        user_message = data.get('message', '').strip()
        input_text = data.get('input_text', '').strip()
        output_text = data.get('output_text', '').strip()
        
        if not user_message:
            return jsonify({'error': 'Message is required'}), 400
        
        # Build context-aware prompt
        system_prompt = """You are a helpful AI assistant integrated into a text processing application. You can help users with any questions, generate content, or discuss their work.

CONTEXT AWARENESS:
- You can see what the user has in their input and output text boxes
- Help them work with, analyze, or improve their content
- Generate new content when requested
- Answer questions about any topic

IMPORTANT: When you provide responses that the user might want to process further (like generated content, essays, problems, etc.), make sure to format them clearly and mention that they can send your response to the input box for further processing."""

        # Add context if there's content in the boxes
        context_info = []
        if input_text:
            context_info.append(f"INPUT BOX CONTENT:\n{input_text[:1000]}{'...' if len(input_text) > 1000 else ''}")
        if output_text:
            context_info.append(f"OUTPUT BOX CONTENT:\n{output_text[:1000]}{'...' if len(output_text) > 1000 else ''}")
        
        if context_info:
            system_prompt += f"\n\nCURRENT CONTEXT:\n" + "\n\n".join(context_info)
        
        # Use the existing chat functionality from ai_processor
        from ai_processor import chat_with_ai as ai_chat
        
        # Combine system prompt and user message
        full_message = f"{system_prompt}\n\nUser: {user_message}"
        
        response = ai_chat(full_message, "")
        
        # ai_chat returns just the response string, not a tuple
        if response and not response.startswith("Error"):
            success = True
        else:
            success = False
            if not response:
                response = "Failed to get response from AI"
        
        if success:
            return jsonify({
                'success': True,
                'response': response
            })
        else:
            return jsonify({'error': f'Chat failed: {response}'}), 500
            
    except Exception as e:
        app.logger.error(f"Error in chat: {str(e)}")
        return jsonify({'error': f'Chat error: {str(e)}'}), 500

@app.route('/homework_direct', methods=['POST'])
def homework_direct():
    """
    Process text directly as homework with LLM, bypassing all rewrite functionality.
    This endpoint sends input text straight to LLM with homework-specific prompting.
    """
    try:
        data = request.get_json()
        text = data.get('text', '').strip()
        
        if not text:
            return jsonify({'error': 'No text provided'}), 400
        
        # Homework-specific prompt that handles ALL academic subjects
        homework_prompt = f"""You are an expert academic tutor specializing in all subjects including mathematics, philosophy, literature, history, science, essays, and any academic discipline. Complete the following assignment with thorough, detailed work:

{text}

Instructions:
1. Identify the subject area and type of assignment (math problem, essay question, philosophy prompt, etc.)
2. Provide complete, detailed responses appropriate to the academic level
3. For mathematics: Show ALL steps with proper LaTeX notation (use \\frac{{a}}{{b}}, x^{{2}}, \\sqrt{{x}}, etc.)
4. For essays/philosophy: Provide structured arguments with clear reasoning and examples
5. For literature: Include textual analysis and proper citations when relevant
6. For science: Explain concepts thoroughly with examples and applications
7. For any subject: Use proper academic formatting and terminology
8. Show your complete work and reasoning process
9. Use LaTeX notation for any mathematical expressions: $...$ for inline, $$...$$ for display

Complete this assignment thoroughly and professionally:"""

        # Use OpenAI directly for mathematical content
        import openai
        
        openai_key = os.environ.get('OPENAI_API_KEY')
        if not openai_key:
            return jsonify({'error': 'OpenAI API key not configured'}), 500
            
        openai_client = openai.OpenAI(api_key=openai_key)
        
        response = openai_client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "user", "content": homework_prompt}
            ],
            max_tokens=2000,
            temperature=0.3
        )
        
        result = response.choices[0].message.content
        
        # Convert dollar signs to "dollars" to prevent LaTeX formatting issues
        result = convert_dollar_signs_to_text(result)
        
        # Remove all markdown formatting from the output
        import re
        # Remove markdown headers (###, ##, #)
        result = re.sub(r'^#{1,6}\s+', '', result, flags=re.MULTILINE)
        # Remove bold formatting (**)
        result = re.sub(r'\*\*(.*?)\*\*', r'\1', result)
        # Remove italic formatting (*)
        result = re.sub(r'\*(.*?)\*', r'\1', result)
        # Remove bullet points (- or *)
        result = re.sub(r'^[\*\-]\s+', '- ', result, flags=re.MULTILINE)
        # Clean up extra whitespace
        result = re.sub(r'\n{3,}', '\n\n', result)
        
        success = True
        
        if success:
            return jsonify({'result': result})
        else:
            return jsonify({'error': result}), 500
        
    except Exception as e:
        logger.error(f"Homework direct processing error: {str(e)}")
        return jsonify({'error': f'Processing failed: {str(e)}'}), 500

@app.route('/style_rewrite_passthrough', methods=['POST'])
def style_rewrite_passthrough_route():
    """
    Process a style rewrite request using the pure pass-through mechanism.
    This endpoint accepts a style sample and target text, and passes them directly
    to the style_rewrite_passthrough module without any modifications.
    
    Request JSON parameters:
    - style_sample: Text representing the user's writing style
    - target_text: Text to be rewritten in the user's style
    
    Returns:
    - JSON response with the rewritten text
    """
    try:
        data = request.get_json()
        style_sample = data.get('style_sample', '')
        target_text = data.get('target_text', '')
        
        if not style_sample or not target_text:
            return jsonify({'error': 'Both style sample and target text are required'}), 400
            
        # Log length of inputs for debugging
        logger.debug(f"Style sample length: {len(style_sample)} chars")
        logger.debug(f"Target text length: {len(target_text)} chars")
        
        # Call the passthrough style rewrite function
        result = process_style_rewrite(style_sample, target_text)
        
        return jsonify({
            'success': True,
            'result': result
        })
        
    except Exception as e:
        logger.error(f"Error in style rewrite passthrough: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/comprehensive_search', methods=['POST'])
def comprehensive_search():
    """
    Perform comprehensive search using Google CSE and multi-AI research.
    Automatically extracts search terms from text content if no query provided.
    """
    try:
        data = request.json
        query = data.get('query', '').strip()
        text_content = data.get('text_content', '').strip()
        
        # Import the comprehensive search module
        from comprehensive_search import perform_comprehensive_search
        
        # Perform the search
        results = perform_comprehensive_search(query=query if query else None, text_content=text_content)
        
        return jsonify({
            'success': True,
            'results': results
        })
        
    except Exception as e:
        logger.error(f"Comprehensive search error: {e}")
        return jsonify({
            'success': False,
            'error': f'Search failed: {str(e)}'
        }), 500